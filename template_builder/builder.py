"""Programmatic DOCX template builder."""

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.core.config import get_settings
from app.core.constants import DocumentType
from app.document_engine.styles import configure_page_setup


class TemplateBuilder:
    """Generate base Word templates from JSON-like configuration."""

    def __init__(self) -> None:
        self.settings = get_settings()

    def _add_styled_paragraph(self, doc: Document, text: str, bold: bool = False) -> None:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.name = self.settings.default_font
        run.font.size = Pt(self.settings.default_font_size)

    def build_from_config(self, config: dict[str, Any], output_path: Path) -> Path:
        doc = Document()
        for section in doc.sections:
            configure_page_setup(section, DocumentType.SOP)

        for block in config.get("sections", []):
            self._render_block(doc, block)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def _render_block(self, doc: Document, block: dict) -> None:
        block_type = block.get("type", "paragraph")

        if block_type == "heading":
            text = block.get("text", "")
            self._add_styled_paragraph(doc, text, bold=True)
        elif block_type == "paragraph":
            self._add_styled_paragraph(doc, block.get("text", ""))
        elif block_type == "jinja_loop":
            loop_var = block.get("loop_var", "tests")
            template_lines = block.get("template_lines", [])
            for line in template_lines:
                self._add_styled_paragraph(doc, line)
        elif block_type == "table_placeholder":
            self._add_styled_paragraph(doc, block.get("text", "{{ table_placeholder }}"))

    def build_moa_template(self, output_path: Path) -> Path:
        config = {
            "document_type": "MOA",
            "sections": [
                {"type": "heading", "text": "{{ product_name }}"},
                {"type": "paragraph", "text": "Reference: {{ reference }}"},
                {"type": "paragraph", "text": "Chemical Formula: {{ chemical_formula }}"},
                {"type": "paragraph", "text": "Molecular Weight: {{ molecular_weight }}"},
                {"type": "paragraph", "text": "MOA No.: {{ moa_no }}"},
                {"type": "paragraph", "text": ""},
                {
                    "type": "jinja_loop",
                    "loop_var": "tests",
                    "template_lines": [
                        "{% for test in tests %}",
                        "{{ test.section_no }} {{ test.name }}",
                        "Procedure:",
                        "{{ test.procedure }}",
                        "Acceptance Criteria: {{ test.acceptance_criteria_display }}",
                        "{% if test.sub_tests %}",
                        "{% for sub in test.sub_tests %}",
                        "{{ sub.section_no }} {{ sub.name }}",
                        "Procedure: {{ sub.procedure }}",
                        "Acceptance Criteria: {{ sub.acceptance_criteria_display }}",
                        "{% endfor %}",
                        "{% endif %}",
                        "",
                        "{% endfor %}",
                    ],
                },
                {
                    "type": "jinja_loop",
                    "loop_var": "additional_tests",
                    "template_lines": [
                        "{% if additional_tests %}",
                        "Additional Tests",
                        "{% for test in additional_tests %}",
                        "{{ test.section_no }} {{ test.name }}",
                        "Procedure: {{ test.procedure }}",
                        "Acceptance Criteria: {{ test.acceptance_criteria_display }}",
                        "{% endfor %}",
                        "{% endif %}",
                    ],
                },
            ],
        }
        return self._build_docxtpl_template(config, output_path)

    def build_protocol_template(self, output_path: Path) -> Path:
        config = {
            "document_type": "PROTOCOL",
            "sections": [
                {"type": "heading", "text": "FINISHED PRODUCT ANALYSIS PROTOCOL"},
                {"type": "paragraph", "text": "Product: {{ product_name }}"},
                {"type": "paragraph", "text": "Protocol No.: {{ protocol_no }}"},
                {"type": "paragraph", "text": "Specification/MOA No.: {{ specification_no }} / {{ moa_no }}"},
                {"type": "paragraph", "text": "Batch No.: {{ batch.batch_no }}"},
                {"type": "paragraph", "text": "Mfg. Date: {{ batch.mfg_date }}"},
                {"type": "paragraph", "text": "Exp. Date: {{ batch.exp_date }}"},
                {"type": "paragraph", "text": "A.R. No.: {{ batch.ar_no }}"},
                {"type": "paragraph", "text": ""},
                {
                    "type": "jinja_loop",
                    "template_lines": [
                        "Sr. No | Tests | Results | Limits",
                        "{% for row in protocol_summary %}",
                        "{{ row.sr_no }}. | {{ row.name }} | {{ row.result }} | {{ row.limit }}",
                        "{% endfor %}",
                    ],
                },
                {"type": "paragraph", "text": ""},
                {
                    "type": "jinja_loop",
                    "template_lines": [
                        "{% for test in all_tests %}",
                        "{{ test.section_no }} {{ test.name }}:",
                        "Procedure:",
                        "{{ test.procedure }}",
                        "Observation: _________________________________________________",
                        "Acceptance Criteria: {{ test.acceptance_criteria_display }}",
                        "Conclusion: Satisfactory / Not satisfactory",
                        "Analyzed By: _______________    Checked By: _______________",
                        "",
                        "{% endfor %}",
                    ],
                },
            ],
        }
        return self._build_docxtpl_template(config, output_path)

    def build_sop_template(self, output_path: Path) -> Path:
        config = {
            "document_type": "SOP",
            "sections": [
                {
                    "type": "jinja_loop",
                    "template_lines": [
                        "{% for section in sop_sections %}",
                        "{{ section.section_no }} {{ section.title }}:",
                        "{% if section.content is string %}",
                        "{{ section.content }}",
                        "{% elif section.content %}",
                        "{% for line in section.content %}",
                        "{{ line }}",
                        "{% endfor %}",
                        "{% endif %}",
                        "{% for sub in section.subsections %}",
                        "{{ sub.section_no }} {{ sub.title }}",
                        "{% if sub.content is string %}{{ sub.content }}{% endif %}",
                        "{% endfor %}",
                        "",
                        "{% endfor %}",
                    ],
                },
            ],
        }
        return self._build_docxtpl_template(config, output_path)

    def build_annexure_template(self, output_path: Path) -> Path:
        config = {
            "document_type": "ANNEXURE",
            "sections": [
                {"type": "heading", "text": "{{ subject }}"},
                {"type": "paragraph", "text": "Annexure No.: {{ document_no }}"},
                {"type": "paragraph", "text": "{{ content }}"},
            ],
        }
        return self._build_docxtpl_template(config, output_path)

    def build_specification_template(self, output_path: Path) -> Path:
        config = {
            "document_type": "SPECIFICATION",
            "sections": [
                {"type": "heading", "text": "{{ product_name }} SPECIFICATION"},
                {"type": "paragraph", "text": "Specification No.: {{ specification_no }}"},
                {"type": "paragraph", "text": "Reference: {{ reference }}"},
                {
                    "type": "jinja_loop",
                    "template_lines": [
                        "Sr. No | Test | Acceptance Criteria",
                        "{% for row in protocol_summary %}",
                        "{{ row.sr_no }}. | {{ row.name }} | {{ row.limit }}",
                        "{% endfor %}",
                    ],
                },
            ],
        }
        return self._build_docxtpl_template(config, output_path)

    def _build_docxtpl_template(self, config: dict, output_path: Path) -> Path:
        doc = Document()
        doc_type_name = config.get("document_type", "SOP")
        try:
            doc_type = DocumentType(doc_type_name)
        except ValueError:
            doc_type = DocumentType.SOP
        for section in doc.sections:
            configure_page_setup(section, doc_type)

        for block in config.get("sections", []):
            if block.get("type") == "jinja_loop":
                for line in block.get("template_lines", []):
                    self._add_styled_paragraph(doc, line)
            else:
                self._render_block(doc, block)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def build_all_base_templates(self, templates_dir: Path) -> dict[str, Path]:
        templates_dir.mkdir(parents=True, exist_ok=True)
        return {
            "base_moa.docx": self.build_moa_template(templates_dir / "base_moa.docx"),
            "base_protocol.docx": self.build_protocol_template(
                templates_dir / "base_protocol.docx"
            ),
            "base_sop.docx": self.build_sop_template(templates_dir / "base_sop.docx"),
            "base_annexure.docx": self.build_annexure_template(
                templates_dir / "base_annexure.docx"
            ),
            "base_specification.docx": self.build_specification_template(
                templates_dir / "base_specification.docx"
            ),
        }
