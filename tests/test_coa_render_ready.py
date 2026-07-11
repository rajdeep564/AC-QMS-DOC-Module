"""Verify COA render path is render-ready — no derivation in layout/context."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.document_engine.renderer import DocumentRenderer
from tests.conftest import FIXTURES_DIR, load_coa_fixture

COA_SOURCE_FILES = (
    Path("app/document_engine/components/coa_layout.py"),
    Path("app/document_engine/coa_context.py"),
    Path("app/schemas/coa_render.py"),
)

FORBIDDEN_PATTERNS = (
    "number_tests",
    "number_sop_sections",
    "acceptance_criteria_display",
    "computeCompliance",
    "formatAcceptance",
    "build_document_context",
    "protocol_summary",
    "to_display_string",
)


@pytest.mark.parametrize("rel_path", COA_SOURCE_FILES, ids=[p.name for p in COA_SOURCE_FILES])
def test_coa_sources_have_no_derivation_imports(rel_path: Path):
    repo_root = Path(__file__).resolve().parents[1]
    text = (repo_root / rel_path).read_text(encoding="utf-8")
    for pattern in FORBIDDEN_PATTERNS:
        assert pattern not in text, f"{rel_path}: must not reference {pattern}"


def test_coa_rendered_text_traces_to_fixture():
    payload = load_coa_fixture("glycine_coa_gcn010226.json")
    renderer = DocumentRenderer()
    path = renderer.render_coa(payload, "test_coa_trace.docx")
    doc = Document(path)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tables_text.append(cell.text)
    all_text = body_text + "\n" + "\n".join(tables_text)

    assert payload.compliance_remark in all_text
    first_row = payload.coa_results[0]
    assert first_row.test_name in all_text
    assert first_row.result in all_text
    assert first_row.conclusion in all_text
