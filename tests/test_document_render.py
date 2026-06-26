"""Smoke tests for programmatic document layouts."""

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from app.core.constants import DocumentType
from app.document_engine.renderer import DocumentRenderer
from app.schemas.product import ProductConfig


@pytest.fixture
def approval():
    return {
        "prepared_by": {"name": "A", "designation": "QA", "date": "01 JAN 2025"},
        "checked_by": {"name": "B", "designation": "QA", "date": "01 JAN 2025"},
        "approved_by": {"name": "C", "designation": "QA Head", "date": "01 JAN 2025"},
    }


@pytest.fixture
def common_kwargs(approval):
    return {
        "revision_no": "01",
        "effective_date": date(2025, 1, 1),
        "review_date": date(2028, 1, 1),
        "superseded_revision": "New",
        "approval": approval,
    }


def _load_config(name: str) -> ProductConfig:
    path = Path(__file__).resolve().parents[1] / "config" / "products" / name
    return ProductConfig.model_validate_json(path.read_text(encoding="utf-8"))


@pytest.mark.parametrize(
    "doc_type,config_name,filename,extra",
    [
        (DocumentType.MOA, "glycine_ip.json", "test_moa.docx", {"document_no": "MOA/FG00038/01"}),
        (DocumentType.PROTOCOL, "glycine_ip.json", "test_protocol.docx", {"document_no": "PROT/FG00038/01", "batch": {"batch_no": "B1"}}),
        (DocumentType.SPECIFICATION, "glycine_ip.json", "test_spec.docx", {"document_no": "SPEC/FG00038/01"}),
        (DocumentType.SOP, "sop_on_sop.json", "test_sop.docx", {"document_no": "QA 01", "subject": "SOP ON SOP"}),
    ],
)
def test_programmatic_render(common_kwargs, doc_type, config_name, filename, extra):
    renderer = DocumentRenderer()
    config = _load_config(config_name)
    path = renderer.render(doc_type, config, filename, **common_kwargs, **extra)
    doc = Document(path)
    assert len(doc.sections) >= 1
    assert doc.sections[0].header.tables
    assert doc.sections[0].footer.tables
