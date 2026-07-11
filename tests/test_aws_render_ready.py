"""Verify AWS render path is render-ready — no derivation in layout/context."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.document_engine.renderer import DocumentRenderer
from tests.conftest import load_aws_fixture

AWS_SOURCE_FILES = (
    Path("app/document_engine/components/aws_layout.py"),
    Path("app/document_engine/aws_context.py"),
    Path("app/schemas/aws_render.py"),
)

FORBIDDEN_PATTERNS = (
    "number_tests",
    "number_sop_sections",
    "acceptance_criteria_display",
    "computeCompliance",
    "formatAcceptance",
    "build_document_context",
    "apply_protocol_layout",
    "protocol_summary",
    "to_display_string",
    "OBSERVATION_LINE",
    "ANALYZED_LINE",
)


@pytest.mark.parametrize("rel_path", AWS_SOURCE_FILES, ids=[p.name for p in AWS_SOURCE_FILES])
def test_aws_sources_have_no_derivation_imports(rel_path: Path):
    repo_root = Path(__file__).resolve().parents[1]
    text = (repo_root / rel_path).read_text(encoding="utf-8")
    for pattern in FORBIDDEN_PATTERNS:
        assert pattern not in text, f"{rel_path}: must not reference {pattern}"


def test_aws_rendered_text_traces_to_fixture():
    payload = load_aws_fixture("glycine_aws_gcn010226.json")
    renderer = DocumentRenderer()
    path = renderer.render_aws(payload, "test_aws_trace.docx")
    doc = Document(path)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    tables_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tables_text.append(cell.text)
    all_text = body_text + "\n" + "\n".join(tables_text)

    oos_section = next(s for s in payload.sections if s.is_oos)
    assert oos_section.oos_ack_comment in all_text
    assert oos_section.result_display in all_text
    assert oos_section.analyst.name in all_text

    expiry_section = next(s for s in payload.sections if s.instrument_expired_ack)
    assert expiry_section.expiry_ack_comment in all_text
