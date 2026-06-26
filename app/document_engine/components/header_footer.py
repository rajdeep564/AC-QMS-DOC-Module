"""Reusable header, footer, and revision history components."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.core.config import get_settings
from app.core.constants import DocumentType
from app.document_engine.components.qa_layout_common import (
    add_page_number_field,
    add_restricted_circulation_line,
    build_qa_approval_footer,
)
from app.document_engine.styles import configure_page_setup, set_cell_text


def build_header_table(
    header,
    context: dict,
) -> None:
    """
    Universal header component per SOP Annexure-I:
    Company | Document Type
    Department | Page No.
    Document No. | Effective Date
    Subject | Review Date
    | Superseded
    """
    settings = get_settings()
    table = header.add_table(rows=5, cols=2, width=Inches(6.5))
    table.autofit = True

    company = context.get("company_name", settings.company_name)
    doc_type = context.get("document_type_label", context.get("document_type", ""))
    department = context.get("department", "QUALITY ASSURANCE")
    document_no = context.get("document_no", "")
    effective_date = context.get("effective_date", "")
    review_date = context.get("review_date", "")
    superseded = context.get("superseded_revision", "")
    subject = context.get("subject", context.get("product_name", ""))

    rows_data = [
        (company, doc_type, True, False),
        (f"DEPARTMENT\n{department}", None, False, False),
        (
            f"{context.get('document_no_label', 'DOCUMENT NO.')}\n{document_no}",
            f"EFFECTIVE DATE\n{effective_date}",
            False,
            False,
        ),
        (f"SUBJECT\n{subject}", f"REVIEW DATE\n{review_date}", False, False),
        ("", f"SUPERSEDED\n{superseded}", False, False),
    ]

    for i, (left, right, left_bold, right_bold) in enumerate(rows_data):
        set_cell_text(table.rows[i].cells[0], left, bold=left_bold)
        if i == 1:
            page_cell = table.rows[i].cells[1]
            page_cell.text = ""
            p = page_cell.paragraphs[0]
            label_run = p.add_run("PAGE NO.\n")
            label_run.bold = right_bold
            label_run.font.name = settings.default_font
            label_run.font.size = Pt(settings.default_font_size)
            add_page_number_field(p)
        elif right is not None:
            set_cell_text(table.rows[i].cells[1], right, bold=right_bold)

    logo_path = context.get("logo_path")
    if logo_path and Path(logo_path).exists():
        try:
            cell = table.rows[0].cells[0]
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(str(logo_path), width=Inches(1.2))
        except Exception:
            pass


def build_footer_table(footer, context: dict) -> None:
    """Universal footer: PREPARED BY | CHECKED BY | APPROVED BY + restricted text."""
    build_qa_approval_footer(footer, context)
    add_restricted_circulation_line(footer)


def add_revision_history_block(doc: Document, entries: list[dict]) -> None:
    """Reusable revision history table."""
    settings = get_settings()
    p = doc.add_paragraph()
    run = p.add_run("Revision History")
    run.bold = True
    run.font.name = settings.default_font
    run.font.size = Pt(settings.default_font_size)

    table = doc.add_table(rows=1 + len(entries), cols=4)
    table.style = "Table Grid"
    headers = ["Document No.", "Revision Made", "Change Control No.", "Effective Date"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)

    for row_idx, entry in enumerate(entries):
        values = [
            entry.get("document_no", ""),
            entry.get("revision_made", ""),
            entry.get("change_control_no", ""),
            str(entry.get("effective_date", "")),
        ]
        for col_idx, val in enumerate(values):
            set_cell_text(table.rows[row_idx + 1].cells[col_idx], val)

    doc.add_paragraph()


def apply_header_footer_to_document(doc: Document, context: dict, document_type: DocumentType | None = None) -> None:
    doc_type = document_type or DocumentType(context.get("document_type", DocumentType.SOP.value))
    for section in doc.sections:
        configure_page_setup(section, doc_type)
        header = section.header
        header.is_linked_to_previous = False
        for p in list(header.paragraphs):
            p._element.getparent().remove(p._element)
        build_header_table(header, context)

        footer = section.footer
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p._element.getparent().remove(p._element)
        build_footer_table(footer, context)
