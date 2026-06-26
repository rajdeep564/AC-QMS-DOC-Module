"""Product Specification layout — exact replica of Spec Glycine IP.docx (Rev 02)."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.core.constants import DocumentType
from app.document_engine.components.qa_layout_common import (
    build_qa_approval_footer,
    build_revision_history_table,
    build_six_col_metadata_header,
    style_qa_cell,
    wire_repeating_header_footer,
)
from docx.table import _Cell as _DocxCell

from app.document_engine.styles import (
    SPEC_FOOTER_GRID_COLS,
    SPEC_FOOTER_TABLE_INDENT_DXA,
    SPEC_FOOTER_TABLE_WIDTH_DXA,
    SPEC_HEADER_GRID_COLS,
    SPEC_HEADER_TABLE_INDENT_DXA,
    SPEC_HEADER_TABLE_WIDTH_DXA,
    SPEC_MICRO_GRID_COLS,
    SPEC_PARAM_GRID_COLS,
    SPEC_PRODUCT_GRID_COLS,
    SPEC_REVISION_GRID_COLS,
    SPEC_REVISION_TABLE_INDENT_DXA,
    SPEC_REVISION_TABLE_WIDTH_DXA,
    SPEC_TABLE_WIDTH_DXA,
    clear_document_body,
    clear_row_height,
    configure_page_setup,
    finalize_protocol_table,
    merge_row_cells,
    set_table_jc,
    setup_protocol_table,
    style_run,
    set_cell_valign,
)


def _pcell(row, col_idx: int):
    """Return the physical TC for a column, bypassing python-docx vMerge resolution.

    row.cells[col] returns the restart cell for continuation rows; this always
    returns the actual TC element that lives in this row's XML.
    """
    return _DocxCell(row._tr.tc_lst[col_idx], row)

# Elemental sub_tests placed in Table 2 (4-col); remainder go in Table 3 (3-col)
_ELEMENTAL_TABLE2_COUNT = 4  # Lead, Arsenic, Mercury, Cadmium in table 2


def _set_cell_vmerge(cell, *, restart: bool = False) -> None:
    """Add vMerge to cell tcPr. restart=True starts merge group; False continues it."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:vMerge"))
    if existing is not None:
        tc_pr.remove(existing)
    vm = OxmlElement("w:vMerge")
    if restart:
        vm.set(qn("w:val"), "restart")
    v_align = tc_pr.find(qn("w:vAlign"))
    if v_align is not None:
        v_align.addprevious(vm)
    else:
        tc_pr.append(vm)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _make_spec_table(doc: Document, n_rows: int, n_cols: int, grid: list) -> object:
    """Create a spec body table: correct width, borders, jc=center."""
    table = doc.add_table(rows=n_rows, cols=n_cols)
    setup_protocol_table(table, grid, width_dxa=SPEC_TABLE_WIDTH_DXA, borders=True)
    set_table_jc(table, "center")
    return table


def _sc(cell, text: str, bold: bool = False,
        align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    """Style a spec body cell (1.5x line spacing via style_qa_cell)."""
    style_qa_cell(cell, text, bold=bold, align=align)


def _sc_numbered(cell, number: int, text: str) -> None:
    """Format a cell with numbered list style matching reference (N. text, bold, hanging indent)."""
    _sc(cell, f"{number}. {text}", bold=True)
    # Apply hanging indent to match reference numPr formatting
    p = cell.paragraphs[0]
    pPr = p._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is not None:
        pPr.remove(ind)
    ind_el = OxmlElement("w:ind")
    ind_el.set(qn("w:left"), "372")
    ind_el.set(qn("w:hanging"), "372")
    pPr.append(ind_el)


def build_spec_header(header, context: dict) -> None:
    mw = context.get("molecular_weight", "")
    if mw and "g/mole" not in str(mw):
        mw = f"{mw} g/mole"
    metadata_rows = [
        ("Name of Material", context.get("product_name", ""), "Specification No.", context.get("specification_no", "")),
        ("Revision No.", context.get("revision_no", "01"), "Page No.", None),
        ("Chemical Formula", context.get("chemical_formula", ""), "Effective Date", str(context.get("effective_date", ""))),
        ("Molecular Weight", mw, "Review Date", str(context.get("review_date", ""))),
        ("Reference", context.get("reference", ""), "Supersedes", context.get("superseded_revision", "")),
    ]
    build_six_col_metadata_header(
        header, context,
        title="FINISHED PRODUCT SPECIFICATION",
        grid_cols=SPEC_HEADER_GRID_COLS,
        table_width_dxa=SPEC_HEADER_TABLE_WIDTH_DXA,
        indent_dxa=SPEC_HEADER_TABLE_INDENT_DXA,
        metadata_rows=metadata_rows,
        formula_value_rows={3},
    )


def build_spec_footer(footer, context: dict) -> None:
    build_qa_approval_footer(
        footer, context,
        grid_cols=SPEC_FOOTER_GRID_COLS,
        table_width_dxa=SPEC_FOOTER_TABLE_WIDTH_DXA,
        indent_dxa=SPEC_FOOTER_TABLE_INDENT_DXA,
    )


# ---------------------------------------------------------------------------
# TABLE 1  Product info + Visual Inspection + Testing Parameters (page 1)
# Grid: [2263, 1701, 1985, 5206]
# Row counter starts at 1, continues into Table 3 for Retained Sample rows.
# ---------------------------------------------------------------------------

def _build_table1(doc: Document, context: dict) -> int:
    """Build Table 1. Returns the next row number for continuation in Table 3."""
    meta = context.get("metadata") or {}
    tests = context.get("tests") or []
    grid = SPEC_PRODUCT_GRID_COLS  # [2263, 1701, 1985, 5206]

    # Row types
    PROD   = "prod"    # merge(0,1) numbered label | merge(2,3) value
    PACK_P = "pack_p"  # col0 vMerge restart numbered "Packing" | col1 "Primary" | merge(2,3) value
    PACK_S = "pack_s"  # col0 vMerge cont   | col1 "Secondary" | merge(2,3) value
    SECT   = "sect"    # col0 vMerge restart numbered label | merge(1,2) "Parameters" | col3 "AC"
    TEST   = "test"    # col0 vMerge cont   | merge(1,2) name | col3 AC
    IDENT  = "ident"   # col0 vMerge cont   | col1 "Identification" | col2 "A) By IR" | col3 AC

    rows = []
    n = [0]  # mutable row counter

    def _prod(label, value):
        n[0] += 1
        rows.append((PROD, label, value, n[0]))

    def _pack_p(value):
        n[0] += 1
        rows.append((PACK_P, "Packing", value, n[0]))

    def _pack_s(value):
        rows.append((PACK_S, "", value, 0))

    def _sect(label):
        n[0] += 1
        rows.append((SECT, label, "", n[0]))

    def _test(name, ac):
        rows.append((TEST, name, ac, 0))

    # Product display name (from metadata if available, else product_name)
    product_display = meta.get("product_display_name") or context.get("product_name", "")

    _prod("Product Name",             product_display)
    _prod("Technical Name",           meta.get("technical_name", ""))
    _prod("Composition",              meta.get("composition", ""))
    _prod("Origin / Source",          meta.get("origin", ""))
    _prod("Method of Production",     meta.get("method_of_production", ""))

    if meta.get("packing_primary") or meta.get("packing_secondary"):
        _pack_p(meta.get("packing_primary", ""))
        _pack_s(meta.get("packing_secondary", ""))

    for label, key in [
        ("Shelf Life of Raw Materials",   "shelf_life"),
        ("Storage Conditions",            "storage_conditions"),
        ("Special Handling Requirements", "special_handling"),
        ("Sampling Plan",                 "sampling_plan"),
    ]:
        val = meta.get(key, "")
        if val:
            _prod(label, val)

    # Visual Inspection Parameters (tests[0] = Description)
    if len(tests) >= 1:
        _sect("Visual Inspection Parameters")
        _test(tests[0].get("name", ""),
              tests[0].get("acceptance_criteria_display", ""))

    # Testing Parameters (tests[1]=Solubility, tests[2]=Identification with sub_tests)
    if len(tests) >= 2:
        _sect("Testing Parameters")
        _test(tests[1].get("name", ""),
              tests[1].get("acceptance_criteria_display", ""))
        if len(tests) >= 3:
            subs = tests[2].get("sub_tests") or []
            if subs:
                # Identification row: col1="Identification", col2="A) By IR", col3=AC
                rows.append((IDENT,
                              "Identification",
                              "A) By IR",
                              subs[0].get("acceptance_criteria_display", "")))

    table = _make_spec_table(doc, len(rows), 4, grid)

    for ri, spec in enumerate(rows):
        row = table.rows[ri]
        rtype = spec[0]

        if rtype == PROD:
            _, label, value, num = spec
            _sc_numbered(row.cells[0], num, label)
            merge_row_cells(row, 0, 1)
            _sc(row.cells[2], value)
            merge_row_cells(row, 2, 3)

        elif rtype == PACK_P:
            _, label, value, num = spec
            _sc_numbered(row.cells[0], num, label)
            _set_cell_vmerge(row.cells[0], restart=True)
            _sc(row.cells[1], "Primary", bold=True)
            _sc(row.cells[2], value)
            merge_row_cells(row, 2, 3)

        elif rtype == PACK_S:
            _, _, value, _ = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], "Secondary", bold=True)
            _sc(row.cells[2], value)
            merge_row_cells(row, 2, 3)

        elif rtype == SECT:
            _, label, _, num = spec
            _sc_numbered(row.cells[0], num, label)
            _set_cell_vmerge(row.cells[0], restart=True)
            _sc(row.cells[1], "Parameters", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_row_cells(row, 1, 2)
            _sc(row.cells[3], "Acceptance Criteria", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        elif rtype == TEST:
            _, name, ac, _ = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], name)
            merge_row_cells(row, 1, 2)
            _sc(row.cells[3], ac)

        elif rtype == IDENT:
            _, parent_name, sub_name, ac = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], parent_name)
            _sc(row.cells[2], sub_name)
            _sc(row.cells[3], ac)

        clear_row_height(row)

    finalize_protocol_table(table, grid)
    return n[0]  # return last used row number


# ---------------------------------------------------------------------------
# TABLE 2  Analytical tests + Additional Tests (incl. first 4 elemental)
# Grid: [2263, 2152, 1534, 5206]
# ---------------------------------------------------------------------------

def _build_table2(doc: Document, context: dict) -> None:
    tests      = context.get("tests") or []
    additional = context.get("additional_tests") or []
    grid = SPEC_PARAM_GRID_COLS  # [2263, 2152, 1534, 5206]

    # Row types
    VHDR  = "vhdr"   # col0 vMerge restart | merge(1,2) "Parameters" | col3 "AC"
    TEST  = "test"   # col0 vMerge cont    | merge(1,2) name | col3 AC
    BTEST = "btest"  # col0 vMerge cont    | col1 EMPTY | col2 name (centered) | col3 AC
    AHDR  = "ahdr"   # col0 vMerge restart | merge(1,2,3) "Additional Tests..."
    ATEST = "atest"  # col0 vMerge cont    | merge(1,2) name | col3 AC
    EHDR  = "ehdr"   # col0 vMerge cont    | merge(1,2,3) "Elemental impurities..."
    ETEST = "etest"  # col0 vMerge cont    | merge(1,2) name | col3 AC

    rows = []
    rows.append((VHDR, "", ""))

    # tests[2] = Identification; sub_tests[1] = "By Chemical test"
    if len(tests) >= 3:
        subs = tests[2].get("sub_tests") or []
        if len(subs) >= 2:
            rows.append((BTEST, f"B) {subs[1].get('name', '')}",
                         subs[1].get("acceptance_criteria_display", "")))

    for t in tests[3:]:
        rows.append((TEST, t.get("name", ""), t.get("acceptance_criteria_display", "")))

    rows.append((AHDR,
                 "Additional Tests (on customer request, this is not part of typical specification)",
                 ""))

    for t in additional[:2]:
        rows.append((ATEST, t.get("name", ""), t.get("acceptance_criteria_display", "")))

    if len(additional) >= 3:
        elmt = additional[2]
        rows.append((EHDR, elmt.get("name", ""), ""))
        for sub in (elmt.get("sub_tests") or [])[:_ELEMENTAL_TABLE2_COUNT]:
            rows.append((ETEST, sub.get("name", ""), sub.get("acceptance_criteria_display", "")))

    table = _make_spec_table(doc, len(rows), 4, grid)

    for ri, spec in enumerate(rows):
        row = table.rows[ri]
        rtype = spec[0]

        if rtype == VHDR:
            _set_cell_vmerge(row.cells[0], restart=True)
            _sc(row.cells[0], "")
            _sc(row.cells[1], "Parameters", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_row_cells(row, 1, 2)
            _sc(row.cells[3], "Acceptance Criteria", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        elif rtype == BTEST:
            # B) Chemical test: col1 EMPTY, col2 = name (not merged), col3 = AC
            _, name, ac = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], "")
            _sc(row.cells[2], name, align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[3], ac)

        elif rtype in (TEST, ATEST, ETEST):
            _, name, ac = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], name)
            merge_row_cells(row, 1, 2)
            _sc(row.cells[3], ac)

        elif rtype == AHDR:
            _, label, _ = spec
            _set_cell_vmerge(row.cells[0], restart=True)
            _sc(row.cells[0], "")
            _sc(row.cells[1], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_row_cells(row, 1, 3)

        elif rtype == EHDR:
            _, label, _ = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_row_cells(row, 1, 3)

        clear_row_height(row)

    finalize_protocol_table(table, grid)


# ---------------------------------------------------------------------------
# TABLE 3  Remaining elemental sub_tests + Microbiological + Retained Sample
# Grid: [2263, 3686, 5206]
# ---------------------------------------------------------------------------

def _build_table3(doc: Document, context: dict, start_num: int = 12) -> None:
    additional  = context.get("additional_tests") or []
    micro_tests = context.get("microbiological_tests") or []
    meta        = context.get("metadata") or {}
    grid = SPEC_MICRO_GRID_COLS  # [2263, 3686, 5206]

    # Row types
    VHDR     = "vhdr"     # col0 vMerge restart | col1 "Parameters" | col2 "AC"
    ETEST    = "etest"    # col0 vMerge cont    | col1 name | col2 AC
    MHDR     = "mhdr"     # col0 vMerge cont    | merge(1,2) "Microbiological parameters"
    MTEST    = "mtest"    # col0 vMerge cont    | col1 name | col2 AC
    RETAINED = "retained" # col0+1 merged numbered label | col2 value (no vMerge)

    rows = []
    rows.append((VHDR, "", ""))

    if len(additional) >= 3:
        for sub in (additional[2].get("sub_tests") or [])[_ELEMENTAL_TABLE2_COUNT:]:
            rows.append((ETEST, sub.get("name", ""), sub.get("acceptance_criteria_display", "")))

    rows.append((MHDR, "Microbiological parameters", ""))
    for mt in micro_tests:
        rows.append((MTEST, mt.get("name", ""), mt.get("acceptance_criteria_display", "")))

    # Retained Sample rows (numbered, continuing from Table 1 counter)
    n = start_num
    if meta.get("retained_sample_quantity"):
        n += 1
        rows.append((RETAINED, "Retained Sample Quantity", meta["retained_sample_quantity"], n))
    if meta.get("retained_sample_period"):
        n += 1
        rows.append((RETAINED, "Retained Sample Retention Period", meta["retained_sample_period"], n))

    table = _make_spec_table(doc, len(rows), 3, grid)

    for ri, spec in enumerate(rows):
        row = table.rows[ri]
        rtype = spec[0]

        if rtype == VHDR:
            _set_cell_vmerge(row.cells[0], restart=True)
            _sc(row.cells[0], "")
            _sc(row.cells[1], "Parameters", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            _sc(row.cells[2], "Acceptance Criteria", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        elif rtype in (ETEST, MTEST):
            _, name, ac = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], name)
            _sc(row.cells[2], ac)

        elif rtype == MHDR:
            _, label, _ = spec
            c0 = _pcell(row, 0)
            _set_cell_vmerge(c0, restart=False)
            _sc(c0, "")
            _sc(row.cells[1], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
            merge_row_cells(row, 1, 2)

        elif rtype == RETAINED:
            _, label, value, num = spec
            _sc_numbered(row.cells[0], num, label)
            merge_row_cells(row, 0, 1)
            _sc(row.cells[2], value)

        clear_row_height(row)

    finalize_protocol_table(table, grid)


# ---------------------------------------------------------------------------
# Public entry points
# ---------------------------------------------------------------------------

def build_spec_body(doc: Document, context: dict) -> None:
    clear_document_body(doc)
    last_num = _build_table1(doc, context)
    _add_page_break(doc)
    _build_table2(doc, context)
    _add_page_break(doc)
    _build_table3(doc, context, start_num=last_num)
    build_revision_history_table(
        doc,
        context.get("revision_history") or [],
        grid_cols=SPEC_REVISION_GRID_COLS,
        table_width_dxa=SPEC_REVISION_TABLE_WIDTH_DXA,
        indent_dxa=SPEC_REVISION_TABLE_INDENT_DXA,
    )


def apply_spec_layout(doc: Document, context: dict) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        configure_page_setup(section, DocumentType.SPECIFICATION)
        wire_repeating_header_footer(
            section, build_spec_header, build_spec_footer, context,
            blank_first_header=False,
            blank_first_footer=False,
        )
    build_spec_body(doc, context)
