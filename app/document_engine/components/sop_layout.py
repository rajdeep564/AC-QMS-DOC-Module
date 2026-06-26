"""SOP layout — Annexure-I header per QA 01 SOP ON SOP with QA approval footer."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from app.core.config import get_settings
from app.core.constants import DocumentType
from app.document_engine.components.qa_layout_common import add_page_number_field
from app.document_engine.components.qa_layout_common import (
    add_restricted_circulation_line,
    build_qa_approval_footer,
    build_revision_history_table,
    clear_header_footer_part,
    style_qa_cell,
    wire_repeating_header_footer,
)
from app.document_engine.styles import (
    PROTOCOL_SPACER_LINE_TWIPS,
    clear_document_body,
    clear_row_height,
    configure_page_setup,
    finalize_protocol_table,
    set_cell_text,
    setup_protocol_table,
    style_run,
)

SOP_HEADER_GRID_COLS = [4680, 4680]
SOP_HEADER_TABLE_WIDTH_DXA = 9360
SOP_REVISION_GRID_COLS = [2340, 2340, 2340, 2340]
SOP_REVISION_TABLE_WIDTH_DXA = 9360


def build_sop_header(header, context: dict) -> None:
    """Annexure-I two-column header per QA SOP ON SOP."""
    settings = get_settings()
    table = header.add_table(rows=5, cols=2, width=Inches(6.5))
    setup_protocol_table(
        table,
        SOP_HEADER_GRID_COLS,
        width_dxa=SOP_HEADER_TABLE_WIDTH_DXA,
        grid_style=True,
    )

    company = context.get("company_name", settings.company_name)
    doc_type = context.get("document_type_label", "STANDARD OPERATING PROCEDURE")
    department = context.get("department", "QUALITY ASSURANCE")
    document_no = context.get("document_no", "")
    effective_date = str(context.get("effective_date", ""))
    review_date = str(context.get("review_date", ""))
    superseded = context.get("superseded_revision", "")
    subject = context.get("subject", context.get("product_name", ""))

    rows_data = [
        (company, doc_type, True, False),
        (f"DEPARTMENT\n{department}", None, False, False),
        (f"{context.get('document_no_label', 'SOP NO.')}\n{document_no}", f"EFFECTIVE DATE\n{effective_date}", False, False),
        (f"SUBJECT\n{subject}", f"REVIEW DATE\n{review_date}", False, False),
        ("", f"SUPERSEDED\n{superseded}", False, False),
    ]

    for i, (left, right, left_bold, right_bold) in enumerate(rows_data):
        set_cell_text(table.rows[i].cells[0], left, bold=left_bold, line_twips=PROTOCOL_SPACER_LINE_TWIPS)
        if i == 1:
            page_cell = table.rows[i].cells[1]
            page_cell.text = ""
            p = page_cell.paragraphs[0]
            label_run = p.add_run("PAGE NO.\n")
            label_run.bold = right_bold
            label_run.font.name = settings.default_font
            label_run.font.size = Pt(settings.default_font_size)
            add_page_number_field(p)
        elif right is not None:
            set_cell_text(table.rows[i].cells[1], right, bold=right_bold, line_twips=PROTOCOL_SPACER_LINE_TWIPS)

    finalize_protocol_table(table, SOP_HEADER_GRID_COLS)


def build_sop_footer(footer, context: dict) -> None:
    build_qa_approval_footer(footer, context)
    add_restricted_circulation_line(footer)


def _add_section_block(doc: Document, section: dict) -> None:
    title = section.get("title", "")
    section_no = section.get("section_no", "")
    heading = doc.add_paragraph()
    style_run(heading.add_run(f"{section_no} {title}:".strip()), bold=True)

    content = section.get("content")
    if isinstance(content, str) and content.strip():
        p = doc.add_paragraph()
        style_run(p.add_run(content))
    elif isinstance(content, list):
        for line in content:
            if str(line).strip():
                p = doc.add_paragraph()
                style_run(p.add_run(str(line)))

    for sub in section.get("subsections") or []:
        sub_no = sub.get("section_no", "")
        sub_title = sub.get("title", "")
        p = doc.add_paragraph()
        style_run(p.add_run(f"{sub_no} {sub_title}".strip()), bold=True)
        sub_content = sub.get("content")
        if isinstance(sub_content, str) and sub_content.strip():
            cp = doc.add_paragraph()
            style_run(cp.add_run(sub_content))


def build_sop_body(doc: Document, context: dict) -> None:
    clear_document_body(doc)
    for section in context.get("sop_sections") or []:
        _add_section_block(doc, section)

    build_revision_history_table(
        doc,
        context.get("revision_history") or [],
        grid_cols=SOP_REVISION_GRID_COLS,
        table_width_dxa=SOP_REVISION_TABLE_WIDTH_DXA,
    )


def apply_sop_layout(doc: Document, context: dict) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        configure_page_setup(section, DocumentType.SOP)
        wire_repeating_header_footer(section, build_sop_header, build_sop_footer, context)
    build_sop_body(doc, context)
