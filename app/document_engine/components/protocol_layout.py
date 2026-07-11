"""Finished Product Analysis Protocol layout — matches GLYCINE IP.docx reference structure."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from app.core.constants import DocumentType
from app.document_engine.components.qa_layout_common import (
    add_logo_or_company,
    add_protocol_page_number_field,
    build_revision_history_table,
)
from app.document_engine.styles import (
    BATCH_GRID_COLS,
    BATCH_TABLE_WIDTH_DXA,
    DETAIL_GRID_COLS,
    FOOTER_GRID_COLS,
    FOOTER_TABLE_INDENT_DXA,
    FOOTER_TABLE_WIDTH_DXA,
    HEADER_GRID_COLS,
    HEADER_LOGO_ROW_HEIGHT_TWIPS,
    HEADER_TABLE_INDENT_DXA,
    HEADER_TABLE_WIDTH_DXA,
    DETAIL_COMPACT_LINE_TWIPS,
    PROTOCOL_REVISION_GRID_COLS,
    PROTOCOL_REVISION_TABLE_INDENT_DXA,
    PROTOCOL_REVISION_TABLE_WIDTH_DXA,
    PROTOCOL_SECTION_GAP_TWIPS,
    PROTOCOL_SPACER_LINE_TWIPS,
    SIGNOFF_GRID_COLS,
    SIGNOFF_TABLE_INDENT_DXA,
    SIGNOFF_TABLE_WIDTH_DXA,
    SUMMARY_GRID_COLS,
    SUMMARY_TABLE_WIDTH_DXA,
    add_header_trailing_paragraph,
    add_footer_trailing_paragraph,
    add_protocol_spacer_paragraph,
    add_zero_gap_paragraph,
    clear_document_body,
    clear_row_height,
    configure_page_setup,
    finalize_protocol_table,
    merge_row_cells,
    set_cell_margins_dxa,
    set_cell_no_wrap,
    set_cell_paragraphs,
    set_cell_text,
    set_cell_valign,
    set_paragraph_spacing_twips,
    set_row_height,
    setup_protocol_table,
    style_run,
)

COMPLIANCE_NOTE = (
    "Note: Above mentioned product complies / does not comply as per IP specification."
)

OBSERVATION_LINE = "Observation: ____________________________________________"
ANALYZED_LINE = (
    "Analyzed By: ___________________"
    "                     Checked By: _____________________"
)


def _clear_header_footer_part(part) -> None:
    for p in list(part.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(part.tables):
        t._element.getparent().remove(t._element)


def _style_header_cell(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    set_cell_text(cell, text, bold=bold, align=align, line_twips=PROTOCOL_SPACER_LINE_TWIPS)
    set_cell_valign(cell, "center")


def build_protocol_header(header, context: dict) -> None:
    """
    Bordered 6-column header matching reference image:
    Row0: logo (left) | FINISHED PRODUCT ANALYSIS PROTOCOL (right, merged)
    Rows1-4: Label : Value | Label : Value with narrow colon columns
    """
    table = header.add_table(rows=5, cols=6, width=Inches(7.67))
    # Use batch column grid so metadata rows align with the batch table below.
    setup_protocol_table(
        table,
        BATCH_GRID_COLS,
        width_dxa=BATCH_TABLE_WIDTH_DXA,
        grid_style=True,
        indent_dxa=HEADER_TABLE_INDENT_DXA,
    )

    product_name = context.get("product_name", "")
    protocol_no = context.get("protocol_no") or context.get("document_no", "")
    batch = context.get("batch") or {}
    revision_no = context.get("revision_no", "01")
    superseded = context.get("superseded_revision") or "New"
    effective_date = str(context.get("effective_date", ""))
    review_date = str(context.get("review_date", ""))

    # Row 0 — logo left, title right (tall row)
    row0 = table.rows[0]
    set_row_height(row0, HEADER_LOGO_ROW_HEIGHT_TWIPS)

    logo_cell = row0.cells[0]
    add_logo_or_company(logo_cell, context)

    merge_row_cells(row0, 1, 5)
    title_cell = row0.cells[1]
    title_cell.text = ""
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing_twips(title_p, line=PROTOCOL_SPACER_LINE_TWIPS)
    style_run(title_p.add_run("FINISHED PRODUCT ANALYSIS PROTOCOL"), bold=True)
    set_cell_valign(title_cell, "center")

    row_defs = [
        ("Protocol ID", protocol_no, "Page No.", None),
        ("Name of Product", product_name, "Batch No.", batch.get("batch_no", "")),
        ("Supersedes", superseded, "Revision No.", revision_no),
        ("Effective Date", effective_date, "Review Date", review_date),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(row_defs, start=1):
        row = table.rows[ri]
        _style_header_cell(row.cells[0], l1, bold=True)
        _style_header_cell(row.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
        _style_header_cell(row.cells[2], str(v1 or ""))
        _style_header_cell(row.cells[3], l2, bold=True)
        _style_header_cell(row.cells[4], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
        if l2 == "Page No.":
            row.cells[5].text = ""
            page_p = row.cells[5].paragraphs[0]
            page_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing_twips(page_p, line=PROTOCOL_SPACER_LINE_TWIPS)
            add_protocol_page_number_field(page_p)
            set_cell_valign(row.cells[5], "center")
        else:
            _style_header_cell(row.cells[5], str(v2 or ""))

    finalize_protocol_table(table, BATCH_GRID_COLS)
    add_header_trailing_paragraph(header)


def build_protocol_footer(footer, context: dict) -> None:
    """6-column bordered footer with merged sign/date pairs."""
    approval = context.get("approval") or {}
    prepared = approval.get("prepared_by") or {}
    checked = approval.get("checked_by") or {}
    approved = approval.get("approved_by") or {}

    table = footer.add_table(rows=5, cols=6, width=Inches(7.67))
    setup_protocol_table(
        table,
        FOOTER_GRID_COLS,
        width_dxa=FOOTER_TABLE_WIDTH_DXA,
        grid_style=True,
        indent_dxa=FOOTER_TABLE_INDENT_DXA,
    )

    row0 = table.rows[0]
    for i, label in enumerate(["PREPARED BY", "CHECKED BY", "APPROVED BY"]):
        _style_header_cell(row0.cells[i * 2], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        merge_row_cells(row0, i * 2, i * 2 + 1)

    for col, label in enumerate(["SIGN", "DATE", "SIGN", "DATE", "SIGN", "DATE"]):
        _style_header_cell(table.rows[1].cells[col], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for col in range(6):
        _style_header_cell(table.rows[2].cells[col], "")

    row3 = table.rows[3]
    for i, name in enumerate([prepared.get("name", ""), checked.get("name", ""), approved.get("name", "")]):
        _style_header_cell(row3.cells[i * 2], name, align=WD_ALIGN_PARAGRAPH.CENTER)
        merge_row_cells(row3, i * 2, i * 2 + 1)

    row4 = table.rows[4]
    for i, desig in enumerate(
        [prepared.get("designation", ""), checked.get("designation", ""), approved.get("designation", "")]
    ):
        _style_header_cell(row4.cells[i * 2], desig, align=WD_ALIGN_PARAGRAPH.CENTER)
        merge_row_cells(row4, i * 2, i * 2 + 1)

    finalize_protocol_table(table, FOOTER_GRID_COLS)
    add_footer_trailing_paragraph(footer)


def wire_repeating_protocol_header_footer(section, context: dict) -> None:
    """
    Logo/protocol header and PREPARED BY footer on every A4 page.

    Word uses first-page parts for page 1 and default parts for pages 2+;
    both must be populated (matches GLYCINE IP.docx).
    """
    section.different_first_page_header_footer = True

    for header_part in (section.first_page_header, section.header):
        header_part.is_linked_to_previous = False
        _clear_header_footer_part(header_part)
        build_protocol_header(header_part, context)

    for footer_part in (section.first_page_footer, section.footer):
        footer_part.is_linked_to_previous = False
        _clear_header_footer_part(footer_part)
        build_protocol_footer(footer_part, context)


def _batch_label_row(table, row_idx: int, ll: str, lv: str, rl: str, rv: str) -> None:
    row = table.rows[row_idx]
    _style_header_cell(row.cells[0], ll, bold=True)
    _style_header_cell(row.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(row.cells[2], lv)
    _style_header_cell(row.cells[3], rl, bold=True)
    _style_header_cell(row.cells[4], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(row.cells[5], rv)


def _build_batch_info_table(doc: Document, context: dict) -> None:
    batch = context.get("batch") or {}
    spec_no = context.get("specification_no", "")
    moa_no = context.get("moa_no", "")
    ref = context.get("reference", "")
    ref_parts = []
    if spec_no:
        ref_parts.append(f"{spec_no} ({ref})" if ref else spec_no)
    if moa_no:
        ref_parts.append(f"{moa_no} ({ref})" if ref else moa_no)
    ref_text = "".join(ref_parts)

    table = doc.add_table(rows=6, cols=6)
    setup_protocol_table(
        table, BATCH_GRID_COLS, width_dxa=BATCH_TABLE_WIDTH_DXA, grid_style=True
    )

    rows_data = [
        ("Mfg. Date", str(batch.get("mfg_date", "")), "Exp. Date", str(batch.get("exp_date", ""))),
        ("Test Request Sheet No.", str(batch.get("test_request_no", "")), "A.R. No.", str(batch.get("ar_no", ""))),
        ("Batch Size", str(batch.get("batch_size", "")), "Quantity Sampled", str(batch.get("quantity_sampled", ""))),
        ("Received Date", str(batch.get("received_date", "")), "Testing Date", str(batch.get("testing_date", ""))),
    ]
    for i, data in enumerate(rows_data):
        _batch_label_row(table, i, *data)

    r4 = table.rows[4]
    _style_header_cell(r4.cells[0], "Reference Specification/MOA No.", bold=True)
    _style_header_cell(r4.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(r4.cells[2], ref_text)
    _style_header_cell(r4.cells[3], "Issuance Stamp by QA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_row_cells(r4, 3, 5)

    r5 = table.rows[5]
    _style_header_cell(r5.cells[0], "Completion Date", bold=True)
    _style_header_cell(r5.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(r5.cells[2], str(batch.get("completion_date", "")))
    _style_header_cell(r5.cells[3], "")
    merge_row_cells(r5, 3, 5)

    finalize_protocol_table(table, BATCH_GRID_COLS)
    add_zero_gap_paragraph(doc)


def _build_summary_table(doc: Document, context: dict) -> None:
    """5-column summary matching GLYCINE IP.docx: Sr.No | Tests | (div) | Results | Limits."""
    rows_data = context.get("protocol_summary") or []
    # SUMMARY_GRID_COLS = [989, 2606, 6, 2787, 4587] — col[2] is a 6 dxa visual divider
    table = doc.add_table(rows=1 + len(rows_data), cols=5)
    setup_protocol_table(
        table, SUMMARY_GRID_COLS, width_dxa=SUMMARY_TABLE_WIDTH_DXA, borders=True
    )

    hdr = table.rows[0]
    set_cell_text(hdr.cells[0], "Sr. No", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr.cells[1], "Tests", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr.cells[2], "")  # divider column — always empty
    set_cell_text(hdr.cells[3], "Results", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr.cells[4], "Limits", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for col in (3, 4):
        set_cell_no_wrap(hdr.cells[col])

    for ri, row in enumerate(rows_data):
        tr = table.rows[ri + 1]
        row_type = row.get("row_type", "standard")
        sr = row.get("sr_no", "")
        name = row.get("name", "")
        result = row.get("result", "")
        limit = row.get("limit", "")

        if row_type == "section_label":
            set_cell_text(tr.cells[0], "")
            set_cell_text(tr.cells[1], name, bold=True)
            merge_row_cells(tr, 1, 4)
        elif row_type == "group_header":
            set_cell_text(tr.cells[0], f"{sr}.", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(tr.cells[1], name)
            merge_row_cells(tr, 1, 4)
        elif row_type == "sub_test":
            set_cell_text(tr.cells[0], "")
            set_cell_text(tr.cells[1], name)
            set_cell_text(tr.cells[2], "")  # divider
            set_cell_text(tr.cells[3], result)
            set_cell_text(tr.cells[4], limit)
        else:
            set_cell_text(tr.cells[0], f"{sr}.", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(tr.cells[1], name)
            set_cell_text(tr.cells[2], "")  # divider
            set_cell_text(tr.cells[3], result)
            set_cell_text(tr.cells[4], limit)

    finalize_protocol_table(table, SUMMARY_GRID_COLS)
    for row in table.rows:
        clear_row_height(row)
    add_protocol_spacer_paragraph(
        doc,
        before=PROTOCOL_SECTION_GAP_TWIPS,
        after=PROTOCOL_SECTION_GAP_TWIPS,
        text=COMPLIANCE_NOTE,
    )


def _build_signoff_table(doc: Document) -> None:
    """Analyzed By / Checked By / Approved By — bordered grid like reference Table 2."""
    table = doc.add_table(rows=3, cols=4)
    setup_protocol_table(
        table,
        SIGNOFF_GRID_COLS,
        width_dxa=SIGNOFF_TABLE_WIDTH_DXA,
        grid_style=True,
        indent_dxa=SIGNOFF_TABLE_INDENT_DXA,
    )

    _style_header_cell(table.rows[0].cells[0], "")
    for col, h in enumerate(["Analyzed By", "Checked By", "Approved By"], start=1):
        _style_header_cell(table.rows[0].cells[col], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(table.rows[1].cells[0], "Sign")
    for col in range(1, 4):
        _style_header_cell(table.rows[1].cells[col], "")
    _style_header_cell(table.rows[2].cells[0], "Date")
    for col in range(1, 4):
        _style_header_cell(table.rows[2].cells[col], "")

    finalize_protocol_table(table, SIGNOFF_GRID_COLS)
    add_zero_gap_paragraph(doc)
    add_protocol_spacer_paragraph(doc)


def _detail_content_paragraphs(test: dict) -> list[str]:
    """Build compact procedure cell content — only lines that exist, no extra blank space."""
    procedure = (test.get("procedure") or "").strip()
    limit = test.get("acceptance_criteria_display", "")
    instruments = test.get("instruments") or []

    lines: list[str] = []
    if instruments:
        lines.append("Instrument Used:")
        lines.extend(line for line in instruments if line.strip())
    if procedure:
        lines.append("Procedure:")
        lines.extend(line for line in procedure.split("\n") if line.strip())
    lines.append(OBSERVATION_LINE)
    if limit:
        lines.append(f"Acceptance criteria: {limit}")
    lines.append("Conclusion: Satisfactory / Not satisfactory")
    lines.append(ANALYZED_LINE)
    return lines


def _sub_test_label(sub: dict, index: int) -> str:
    """Map sub-test to A., B., C. labels like the reference document."""
    section_no = sub.get("section_no", "")
    if section_no and "." in section_no:
        last = section_no.split(".")[-1]
        if last.isalpha() and len(last) == 1:
            return f"{last.upper()}."
    return chr(65 + index) + "."


def _new_detail_table(doc: Document, row_count: int):
    table = doc.add_table(rows=row_count, cols=3)
    setup_protocol_table(
        table, DETAIL_GRID_COLS, borders=True, indent_dxa=HEADER_TABLE_INDENT_DXA
    )
    return table


def _write_detail_title_row(row, col0: str, title: str) -> None:
    set_cell_text(row.cells[0], col0, bold=True, line_twips=DETAIL_COMPACT_LINE_TWIPS)
    set_cell_text(row.cells[1], title, line_twips=DETAIL_COMPACT_LINE_TWIPS)
    merge_row_cells(row, 1, 2)
    for cell in row.cells:
        set_cell_margins_dxa(cell)
    clear_row_height(row)


def _write_detail_content_row(row, lines: list[str]) -> None:
    cell = row.cells[0]
    set_cell_paragraphs(cell, lines, compact=True)
    merge_row_cells(row, 0, 2)
    set_cell_valign(cell, "top")
    set_cell_margins_dxa(cell)
    clear_row_height(row)


def _build_test_detail_block(doc: Document, test: dict, *, add_spacer_after: bool) -> None:
    """One or more small tables per test — prevents empty pages from one giant table."""
    section_no = test.get("section_no", "")
    name = test.get("name", "")
    sub_tests = test.get("sub_tests") or []

    if sub_tests:
        table = _new_detail_table(doc, 1)
        _write_detail_title_row(table.rows[0], section_no, f"{name}:")
        finalize_protocol_table(table, DETAIL_GRID_COLS)
        clear_row_height(table.rows[0])

        for si, sub in enumerate(sub_tests):
            table = _new_detail_table(doc, 2)
            _write_detail_title_row(
                table.rows[0],
                _sub_test_label(sub, si),
                f"{sub.get('name', '')}:",
            )
            _write_detail_content_row(table.rows[1], _detail_content_paragraphs(sub))
            finalize_protocol_table(table, DETAIL_GRID_COLS)
            if si < len(sub_tests) - 1 or add_spacer_after:
                add_protocol_spacer_paragraph(doc, before=0, after=0, line=DETAIL_COMPACT_LINE_TWIPS)
    else:
        table = _new_detail_table(doc, 2)
        _write_detail_title_row(table.rows[0], section_no, f"{name}:")
        _write_detail_content_row(table.rows[1], _detail_content_paragraphs(test))
        finalize_protocol_table(table, DETAIL_GRID_COLS)
        if add_spacer_after:
            add_protocol_spacer_paragraph(doc, before=0, after=0, line=DETAIL_COMPACT_LINE_TWIPS)


def _build_detail_tables(doc: Document, context: dict) -> None:
    """Detail procedures as separate tables per test — matches reference pagination."""
    tests = context.get("all_tests") or []
    for i, test in enumerate(tests):
        _build_test_detail_block(doc, test, add_spacer_after=i < len(tests) - 1)


def build_protocol_body(doc: Document, context: dict) -> None:
    clear_document_body(doc)
    _build_batch_info_table(doc, context)
    _build_summary_table(doc, context)
    _build_signoff_table(doc)
    _build_detail_tables(doc, context)
    build_revision_history_table(
        doc,
        context.get("revision_history") or [],
        grid_cols=PROTOCOL_REVISION_GRID_COLS,
        table_width_dxa=PROTOCOL_REVISION_TABLE_WIDTH_DXA,
        indent_dxa=PROTOCOL_REVISION_TABLE_INDENT_DXA,
        label="Protocol No.",
    )


def apply_protocol_layout(doc: Document, context: dict) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        configure_page_setup(section, DocumentType.PROTOCOL)
        wire_repeating_protocol_header_footer(section, context)

    build_protocol_body(doc, context)
