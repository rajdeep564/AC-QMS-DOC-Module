"""Shared QA document layout: footers, revision history, header/footer wiring."""

import re
import struct
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.config import get_settings
from app.core.constants import FOOTER_RESTRICTED_TEXT
from app.document_engine.styles import (
    FOOTER_GRID_COLS,
    FOOTER_TABLE_INDENT_DXA,
    FOOTER_TABLE_WIDTH_DXA,
    HEADER_LOGO_ROW_HEIGHT_TWIPS,
    PROTOCOL_SPACER_LINE_TWIPS,
    add_footer_trailing_paragraph,
    add_header_trailing_paragraph,
    clear_row_height,
    finalize_protocol_table,
    merge_row_cells,
    set_cell_text,
    set_cell_valign,
    set_paragraph_spacing_twips,
    set_row_height,
    setup_protocol_table,
    style_run,
)


def _logo_height_inches(logo_path: str, target_width_inches: float) -> float | None:
    """Return proportional height for logo at target_width, reading actual pixel dims."""
    try:
        path = str(logo_path)
        with open(path, "rb") as f:
            data = f.read(512)
        if data[:8] == b"\x89PNG\r\n\x1a\n":
            w = struct.unpack(">I", data[16:20])[0]
            h = struct.unpack(">I", data[20:24])[0]
        else:
            # JPEG: scan for SOF marker
            w = h = 0
            with open(path, "rb") as f:
                raw = f.read()
            i = 2
            while i < len(raw) - 3:
                if raw[i] != 0xFF:
                    break
                marker = raw[i + 1]
                if marker in (0xC0, 0xC1, 0xC2, 0xC3):
                    h = struct.unpack(">H", raw[i + 5 : i + 7])[0]
                    w = struct.unpack(">H", raw[i + 7 : i + 9])[0]
                    break
                length = struct.unpack(">H", raw[i + 2 : i + 4])[0]
                i += 2 + length
        if w and h:
            return target_width_inches * h / w
    except Exception:
        pass
    return None


def _append_page_field(run, instruction: str) -> None:
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    run._r.append(instr)

    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_char_sep)
    run._r.append(OxmlElement("w:t"))

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def add_page_number_field(paragraph) -> None:
    """Insert PAGE and NUMPAGES fields for 'Page X of Y'."""
    run = paragraph.add_run("Page ")
    _append_page_field(run, " PAGE ")
    run = paragraph.add_run(" of ")
    _append_page_field(run, " NUMPAGES ")


def add_protocol_page_number_field(paragraph) -> None:
    """Product doc header format: '4 of 39' — matches GLYCINE IP.docx."""
    run = paragraph.add_run()
    _append_page_field(run, " PAGE  \\* Arabic  \\* MERGEFORMAT ")
    run = paragraph.add_run(" of ")
    _append_page_field(run, " NUMPAGES  \\* Arabic  \\* MERGEFORMAT ")


def clear_header_footer_part(part) -> None:
    for p in list(part.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(part.tables):
        t._element.getparent().remove(t._element)


def style_qa_cell(
    cell,
    text: str,
    *,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    set_cell_text(cell, text, bold=bold, align=align, line_twips=PROTOCOL_SPACER_LINE_TWIPS)
    set_cell_valign(cell, "center")


def set_chemical_formula_cell(cell, formula: str) -> None:
    """Render chemical formula with proper subscript numbers (C2H5NO2 → C₂H₅NO₂)."""
    settings = get_settings()
    cell.text = ""
    p = cell.paragraphs[0]
    set_paragraph_spacing_twips(p, line=PROTOCOL_SPACER_LINE_TWIPS)
    for part in re.split(r"(\d+)", formula):
        if not part:
            continue
        run = p.add_run(part)
        run.font.name = settings.default_font
        run.font.size = Pt(settings.default_font_size)
        if part.isdigit():
            rPr = run._r.get_or_add_rPr()
            vert = OxmlElement("w:vertAlign")
            vert.set(qn("w:val"), "subscript")
            rPr.append(vert)
    set_cell_valign(cell, "center")


def build_qa_approval_footer(
    footer,
    context: dict,
    *,
    grid_cols: list[int] | None = None,
    table_width_dxa: int | None = None,
    indent_dxa: int | None = None,
) -> None:
    """6-column PREPARED BY / CHECKED BY / APPROVED BY footer — all QA product docs."""
    approval = context.get("approval") or {}
    prepared = approval.get("prepared_by") or {}
    checked = approval.get("checked_by") or {}
    approved = approval.get("approved_by") or {}

    cols = grid_cols or FOOTER_GRID_COLS
    table = footer.add_table(rows=5, cols=6, width=Inches(7.67))
    setup_protocol_table(
        table,
        cols,
        width_dxa=table_width_dxa or FOOTER_TABLE_WIDTH_DXA,
        grid_style=True,
        indent_dxa=indent_dxa if indent_dxa is not None else FOOTER_TABLE_INDENT_DXA,
    )

    row0 = table.rows[0]
    for i, label in enumerate(["PREPARED BY", "CHECKED BY", "APPROVED BY"]):
        style_qa_cell(row0.cells[i * 2], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        merge_row_cells(row0, i * 2, i * 2 + 1)

    for col, label in enumerate(["SIGN", "DATE", "SIGN", "DATE", "SIGN", "DATE"]):
        style_qa_cell(table.rows[1].cells[col], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for col in range(6):
        style_qa_cell(table.rows[2].cells[col], "")

    row3 = table.rows[3]
    for i, name in enumerate([prepared.get("name", ""), checked.get("name", ""), approved.get("name", "")]):
        style_qa_cell(row3.cells[i * 2], name, align=WD_ALIGN_PARAGRAPH.CENTER)
        merge_row_cells(row3, i * 2, i * 2 + 1)

    row4 = table.rows[4]
    for i, desig in enumerate(
        [prepared.get("designation", ""), checked.get("designation", ""), approved.get("designation", "")]
    ):
        style_qa_cell(row4.cells[i * 2], desig, align=WD_ALIGN_PARAGRAPH.CENTER)
        merge_row_cells(row4, i * 2, i * 2 + 1)

    finalize_protocol_table(table, cols)
    add_footer_trailing_paragraph(footer)


def add_restricted_circulation_line(footer) -> None:
    settings = get_settings()
    p = footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER_RESTRICTED_TEXT)
    run.font.name = settings.default_font
    run.font.size = Pt(settings.default_font_size)
    run.bold = True


def wire_repeating_header_footer(
    section,
    header_builder,
    footer_builder,
    context: dict,
    *,
    blank_first_header: bool = False,
    blank_first_footer: bool = False,
) -> None:
    """Populate header/footer parts on all pages.

    blank_first_header=True: first-page header is intentionally empty (e.g. Spec doc).
    blank_first_footer=True: first-page footer is intentionally empty (e.g. Spec doc).
    """
    section.different_first_page_header_footer = True

    if blank_first_header:
        fp_hdr = section.first_page_header
        fp_hdr.is_linked_to_previous = False
        clear_header_footer_part(fp_hdr)
        fp_hdr.add_paragraph()
    else:
        fp_hdr = section.first_page_header
        fp_hdr.is_linked_to_previous = False
        clear_header_footer_part(fp_hdr)
        header_builder(fp_hdr, context)

    section.header.is_linked_to_previous = False
    clear_header_footer_part(section.header)
    header_builder(section.header, context)

    if blank_first_footer:
        fp_ftr = section.first_page_footer
        fp_ftr.is_linked_to_previous = False
        clear_header_footer_part(fp_ftr)
        fp_ftr.add_paragraph()
    else:
        fp_ftr = section.first_page_footer
        fp_ftr.is_linked_to_previous = False
        clear_header_footer_part(fp_ftr)
        footer_builder(fp_ftr, context)

    section.footer.is_linked_to_previous = False
    clear_header_footer_part(section.footer)
    footer_builder(section.footer, context)


def add_logo_or_company(cell, context: dict, *, logo_width: float = 1.33) -> None:
    settings_company = context.get("company_name", "Aditya Chemicals")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # No line-spacing on image paragraphs — extra line spacing stretches the logo row
    logo_path = context.get("logo_path")
    if logo_path and Path(logo_path).exists():
        try:
            logo_h = _logo_height_inches(logo_path, logo_width)
            if logo_h:
                p.add_run().add_picture(str(logo_path), width=Inches(logo_width), height=Inches(logo_h))
            else:
                p.add_run().add_picture(str(logo_path), width=Inches(logo_width))
        except Exception:
            style_run(p.add_run(settings_company), bold=True)
    else:
        style_run(p.add_run(settings_company), bold=True)
    set_cell_valign(cell, "center")


def add_centered_title(cell, title: str) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing_twips(p, line=PROTOCOL_SPACER_LINE_TWIPS)
    style_run(p.add_run(title), bold=True)
    set_cell_valign(cell, "center")


def build_six_col_metadata_header(
    header,
    context: dict,
    *,
    title: str,
    grid_cols: list[int],
    table_width_dxa: int,
    indent_dxa: int,
    metadata_rows: list[tuple[str, str, str, str | None]],
    formula_value_rows: set[int] | None = None,
):
    """Bordered 6-column header: logo + title row, then label : value pairs.

    formula_value_rows: 1-based row indices (from metadata_rows) where v1 is a
    chemical formula and should be rendered with subscript numbers.
    Returns the header table for post-processing.
    """
    formula_rows = formula_value_rows or set()
    table = header.add_table(rows=1 + len(metadata_rows), cols=6, width=Inches(7.67))
    setup_protocol_table(
        table, grid_cols, width_dxa=table_width_dxa, grid_style=True, indent_dxa=indent_dxa
    )

    row0 = table.rows[0]
    set_row_height(row0, HEADER_LOGO_ROW_HEIGHT_TWIPS)
    add_logo_or_company(row0.cells[0], context)
    merge_row_cells(row0, 1, 5)
    add_centered_title(row0.cells[1], title)

    for ri, (l1, v1, l2, v2) in enumerate(metadata_rows, start=1):
        row = table.rows[ri]
        style_qa_cell(row.cells[0], l1, bold=True)
        style_qa_cell(row.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
        if ri in formula_rows and v1:
            set_chemical_formula_cell(row.cells[2], str(v1))
        else:
            style_qa_cell(row.cells[2], str(v1 or ""))
        style_qa_cell(row.cells[3], l2, bold=True)
        style_qa_cell(row.cells[4], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
        if l2 == "Page No.":
            row.cells[5].text = ""
            page_p = row.cells[5].paragraphs[0]
            set_paragraph_spacing_twips(page_p, line=PROTOCOL_SPACER_LINE_TWIPS)
            add_protocol_page_number_field(page_p)
            set_cell_valign(row.cells[5], "center")
        else:
            style_qa_cell(row.cells[5], str(v2 or ""))

    finalize_protocol_table(table, grid_cols)
    add_header_trailing_paragraph(header)
    return table


def build_revision_history_table(
    doc: Document,
    entries: list[dict],
    *,
    grid_cols: list[int],
    table_width_dxa: int,
    indent_dxa: int = 0,
    label: str = "Specification No.",
) -> None:
    if not entries:
        return

    p = doc.add_paragraph()
    style_run(p.add_run("Revision History:"), bold=True)

    table = doc.add_table(rows=1 + len(entries), cols=4)
    setup_protocol_table(
        table, grid_cols, width_dxa=table_width_dxa, grid_style=True, indent_dxa=indent_dxa
    )

    headers = [label, "Revision Made", "Ref. CC No.", "Effective Date"]
    for i, h in enumerate(headers):
        style_qa_cell(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, entry in enumerate(entries):
        row = table.rows[ri + 1]
        values = [
            str(entry.get("document_no", "")),
            str(entry.get("revision_made", "")),
            str(entry.get("change_control_no", "")),
            str(entry.get("effective_date", "")),
        ]
        for ci, val in enumerate(values):
            style_qa_cell(row.cells[ci], val)

    finalize_protocol_table(table, grid_cols)
    for row in table.rows:
        clear_row_height(row)
