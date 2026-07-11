"""Certificate of Analysis layout — flat results table + compliance verdict."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.constants import DocumentType
from app.document_engine.components.qa_layout_common import (
    build_qa_approval_footer,
    build_revision_history_table,
    build_six_col_metadata_header,
    style_qa_cell,
    wire_repeating_header_footer,
)
from app.document_engine.styles import (
    COA_FOOTER_GRID_COLS,
    COA_FOOTER_TABLE_INDENT_DXA,
    COA_FOOTER_TABLE_WIDTH_DXA,
    COA_HEADER_GRID_COLS,
    COA_HEADER_TABLE_INDENT_DXA,
    COA_HEADER_TABLE_WIDTH_DXA,
    COA_RESULTS_GRID_COLS,
    COA_RESULTS_TABLE_WIDTH_DXA,
    COA_REVISION_GRID_COLS,
    COA_REVISION_TABLE_INDENT_DXA,
    COA_REVISION_TABLE_WIDTH_DXA,
    clear_document_body,
    clear_row_height,
    configure_page_setup,
    finalize_protocol_table,
    set_table_jc,
    setup_protocol_table,
    style_run,
)


def build_coa_header(header, context: dict) -> None:
    product = context.get("product") or {}
    batch = context.get("batch") or {}
    metadata_rows = [
        (
            "Name of Material",
            product.get("product_name", ""),
            "Batch No.",
            batch.get("batch_no", ""),
        ),
        (
            context.get("document_no_label", "COA NO."),
            context.get("document_no", ""),
            "A.R. No.",
            batch.get("arn_no", ""),
        ),
        (
            "Specification No.",
            product.get("specification_no", ""),
            "Revision No.",
            context.get("revision_no", ""),
        ),
        (
            "Effective Date",
            str(context.get("effective_date", "")),
            "Page No.",
            None,
        ),
        (
            "Reference",
            product.get("reference", ""),
            "Review Date",
            str(context.get("review_date", "")),
        ),
    ]
    build_six_col_metadata_header(
        header,
        context,
        title=context.get("document_type_label", "ANALYTICAL REPORT"),
        grid_cols=COA_HEADER_GRID_COLS,
        table_width_dxa=COA_HEADER_TABLE_WIDTH_DXA,
        indent_dxa=COA_HEADER_TABLE_INDENT_DXA,
        metadata_rows=metadata_rows,
    )


def build_coa_footer(footer, context: dict) -> None:
    build_qa_approval_footer(
        footer,
        context,
        grid_cols=COA_FOOTER_GRID_COLS,
        table_width_dxa=COA_FOOTER_TABLE_WIDTH_DXA,
        indent_dxa=COA_FOOTER_TABLE_INDENT_DXA,
    )


def _build_coa_results_table(doc: Document, context: dict) -> None:
    rows_data = context.get("coa_results") or []
    if not rows_data:
        return

    table = doc.add_table(rows=1 + len(rows_data), cols=4)
    setup_protocol_table(
        table,
        COA_RESULTS_GRID_COLS,
        width_dxa=COA_RESULTS_TABLE_WIDTH_DXA,
        borders=True,
    )
    set_table_jc(table, "center")

    headers = ["Test", "Result", "Acceptance Limits", "Conclusion"]
    hdr = table.rows[0]
    for col, label in enumerate(headers):
        style_qa_cell(hdr.cells[col], label, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, row in enumerate(rows_data):
        tr = table.rows[ri + 1]
        style_qa_cell(tr.cells[0], row.get("test_name", ""))
        style_qa_cell(tr.cells[1], row.get("result", ""))
        style_qa_cell(tr.cells[2], row.get("acceptance_limits") or "")
        style_qa_cell(tr.cells[3], row.get("conclusion") or "")

    finalize_protocol_table(table, COA_RESULTS_GRID_COLS)
    for row in table.rows:
        clear_row_height(row)


def _build_coa_verdict(doc: Document, context: dict) -> None:
    remark = context.get("compliance_remark", "")
    if not remark:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run(remark), bold=True)


def build_coa_body(doc: Document, context: dict) -> None:
    clear_document_body(doc)
    _build_coa_results_table(doc, context)
    _build_coa_verdict(doc, context)
    if context.get("revision_history"):
        build_revision_history_table(
            doc,
            context.get("revision_history") or [],
            grid_cols=COA_REVISION_GRID_COLS,
            table_width_dxa=COA_REVISION_TABLE_WIDTH_DXA,
            indent_dxa=COA_REVISION_TABLE_INDENT_DXA,
            label="COA No.",
        )


def apply_coa_layout(doc: Document, context: dict) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        configure_page_setup(section, DocumentType.COA)
        wire_repeating_header_footer(
            section,
            build_coa_header,
            build_coa_footer,
            context,
        )
    build_coa_body(doc, context)
