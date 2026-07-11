"""Analytical Work Sheet layout — filled protocol-style body with render-ready payload."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from app.core.constants import DocumentType
from app.document_engine.components.qa_layout_common import (
    add_logo_or_company,
    add_protocol_page_number_field,
    build_qa_approval_footer,
    build_revision_history_table,
)
from app.document_engine.styles import (
    AWS_BATCH_GRID_COLS,
    AWS_BATCH_TABLE_WIDTH_DXA,
    AWS_DETAIL_GRID_COLS,
    AWS_FOOTER_GRID_COLS,
    AWS_FOOTER_TABLE_INDENT_DXA,
    AWS_FOOTER_TABLE_WIDTH_DXA,
    AWS_HEADER_TABLE_INDENT_DXA,
    AWS_HEADER_TABLE_WIDTH_DXA,
    AWS_REVISION_GRID_COLS,
    AWS_REVISION_TABLE_INDENT_DXA,
    AWS_REVISION_TABLE_WIDTH_DXA,
    AWS_SUMMARY_GRID_COLS,
    AWS_SUMMARY_TABLE_WIDTH_DXA,
    DETAIL_COMPACT_LINE_TWIPS,
    HEADER_LOGO_ROW_HEIGHT_TWIPS,
    PROTOCOL_SECTION_GAP_TWIPS,
    PROTOCOL_SPACER_LINE_TWIPS,
    add_footer_trailing_paragraph,
    add_header_trailing_paragraph,
    add_protocol_spacer_paragraph,
    add_zero_gap_paragraph,
    clear_document_body,
    clear_row_height,
    configure_page_setup,
    finalize_protocol_table,
    merge_row_cells,
    set_cell_margins_dxa,
    set_cell_no_wrap,
    set_cell_paragraphs,
    set_cell_text,
    set_cell_valign,
    set_paragraph_spacing_twips,
    set_row_height,
    setup_protocol_table,
    style_run,
)

DEFAULT_COMPLIANCE_NOTE = (
    "Note: Above mentioned product complies / does not comply as per IP specification."
)


def _clear_header_footer_part(part) -> None:
    for p in list(part.paragraphs):
        p._element.getparent().remove(p._element)
    for t in list(part.tables):
        t._element.getparent().remove(t._element)


def _style_header_cell(cell, text: str, *, bold: bool = False, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    set_cell_text(cell, text, bold=bold, align=align, line_twips=PROTOCOL_SPACER_LINE_TWIPS)
    set_cell_valign(cell, "center")


def build_aws_header(header, context: dict) -> None:
    """Bordered 6-column header — title from document_type_label, AWS No. from payload."""
    table = header.add_table(rows=5, cols=6, width=Inches(7.67))
    setup_protocol_table(
        table,
        AWS_BATCH_GRID_COLS,
        width_dxa=AWS_BATCH_TABLE_WIDTH_DXA,
        grid_style=True,
        indent_dxa=AWS_HEADER_TABLE_INDENT_DXA,
    )

    product_name = context.get("product_name", "")
    aws_no = context.get("document_no", "")
    doc_label = context.get("document_no_label", "AWS NO.")
    title = context.get("document_type_label", "FINISHED PRODUCT ANALYSIS PROTOCOL")
    batch = context.get("batch") or {}
    revision_no = context.get("revision_no", "01")
    superseded = context.get("superseded_revision") or "New"
    effective_date = str(context.get("effective_date", ""))
    review_date = str(context.get("review_date", ""))

    row0 = table.rows[0]
    set_row_height(row0, HEADER_LOGO_ROW_HEIGHT_TWIPS)

    logo_cell = row0.cells[0]
    add_logo_or_company(logo_cell, context)

    merge_row_cells(row0, 1, 5)
    title_cell = row0.cells[1]
    title_cell.text = ""
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing_twips(title_p, line=PROTOCOL_SPACER_LINE_TWIPS)
    style_run(title_p.add_run(title), bold=True)
    set_cell_valign(title_cell, "center")

    row_defs = [
        (doc_label, aws_no, "Page No.", None),
        ("Name of Product", product_name, "Batch No.", batch.get("batch_no", "")),
        ("Supersedes", superseded, "Revision No.", revision_no),
        ("Effective Date", effective_date, "Review Date", review_date),
    ]
    for ri, (l1, v1, l2, v2) in enumerate(row_defs, start=1):
        row = table.rows[ri]
        _style_header_cell(row.cells[0], l1, bold=True)
        _style_header_cell(row.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
        _style_header_cell(row.cells[2], str(v1 or ""))
        _style_header_cell(row.cells[3], l2, bold=True)
        _style_header_cell(row.cells[4], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
        if l2 == "Page No.":
            row.cells[5].text = ""
            page_p = row.cells[5].paragraphs[0]
            page_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_spacing_twips(page_p, line=PROTOCOL_SPACER_LINE_TWIPS)
            add_protocol_page_number_field(page_p)
            set_cell_valign(row.cells[5], "center")
        else:
            _style_header_cell(row.cells[5], str(v2 or ""))

    finalize_protocol_table(table, AWS_BATCH_GRID_COLS)
    add_header_trailing_paragraph(header)


def build_aws_footer(footer, context: dict) -> None:
    build_qa_approval_footer(
        footer,
        context,
        grid_cols=AWS_FOOTER_GRID_COLS,
        table_width_dxa=AWS_FOOTER_TABLE_WIDTH_DXA,
        indent_dxa=AWS_FOOTER_TABLE_INDENT_DXA,
    )
    add_footer_trailing_paragraph(footer)


def wire_repeating_aws_header_footer(section, context: dict) -> None:
    section.different_first_page_header_footer = True

    for header_part in (section.first_page_header, section.header):
        header_part.is_linked_to_previous = False
        _clear_header_footer_part(header_part)
        build_aws_header(header_part, context)

    for footer_part in (section.first_page_footer, section.footer):
        footer_part.is_linked_to_previous = False
        _clear_header_footer_part(footer_part)
        build_aws_footer(footer_part, context)


def _batch_label_row(table, row_idx: int, ll: str, lv: str, rl: str, rv: str) -> None:
    row = table.rows[row_idx]
    _style_header_cell(row.cells[0], ll, bold=True)
    _style_header_cell(row.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(row.cells[2], lv)
    _style_header_cell(row.cells[3], rl, bold=True)
    _style_header_cell(row.cells[4], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(row.cells[5], rv)


def _build_batch_info_table(doc: Document, context: dict) -> None:
    batch = context.get("batch") or {}
    spec_no = context.get("specification_no", "")
    moa_no = context.get("moa_no", "")
    ref = context.get("reference", "")
    ref_parts = []
    if spec_no:
        ref_parts.append(f"{spec_no} ({ref})" if ref else spec_no)
    if moa_no:
        ref_parts.append(f"{moa_no} ({ref})" if ref else moa_no)
    ref_text = "".join(ref_parts)

    table = doc.add_table(rows=6, cols=6)
    setup_protocol_table(
        table, AWS_BATCH_GRID_COLS, width_dxa=AWS_BATCH_TABLE_WIDTH_DXA, grid_style=True
    )

    rows_data = [
        ("Mfg. Date", str(batch.get("mfg_date", "")), "Exp. Date", str(batch.get("exp_date", ""))),
        (
            "Test Request Sheet No.",
            str(batch.get("test_request_no", "")),
            "A.R. No.",
            str(batch.get("ar_no", "")),
        ),
        (
            "Batch Size",
            str(batch.get("batch_size", "")),
            "Quantity Sampled",
            str(batch.get("quantity_sampled", "")),
        ),
        (
            "Received Date",
            str(batch.get("received_date", "")),
            "Testing Date",
            str(batch.get("testing_date", "")),
        ),
    ]
    for i, data in enumerate(rows_data):
        _batch_label_row(table, i, *data)

    r4 = table.rows[4]
    _style_header_cell(r4.cells[0], "Reference Specification/MOA No.", bold=True)
    _style_header_cell(r4.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(r4.cells[2], ref_text)
    _style_header_cell(r4.cells[3], "Issuance Stamp by QA", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    merge_row_cells(r4, 3, 5)

    r5 = table.rows[5]
    _style_header_cell(r5.cells[0], "Completion Date", bold=True)
    _style_header_cell(r5.cells[1], ":", align=WD_ALIGN_PARAGRAPH.CENTER)
    _style_header_cell(r5.cells[2], str(batch.get("completion_date", "")))
    _style_header_cell(r5.cells[3], "")
    merge_row_cells(r5, 3, 5)

    finalize_protocol_table(table, AWS_BATCH_GRID_COLS)
    add_zero_gap_paragraph(doc)


def _project_summary_from_sections(sections: list[dict]) -> list[dict]:
    rows: list[dict] = []
    for i, section in enumerate(sections):
        sr = section.get("section_no") or str(i + 1)
        rows.append(
            {
                "row_type": "standard",
                "sr_no": sr,
                "name": section.get("test_name", ""),
                "result": section.get("result_display", ""),
                "limit": section.get("limits_display", ""),
            }
        )
    return rows


def _build_aws_summary_table(doc: Document, context: dict) -> None:
    rows_data = context.get("summary_rows")
    if rows_data is None:
        rows_data = _project_summary_from_sections(context.get("sections") or [])

    table = doc.add_table(rows=1 + len(rows_data), cols=5)
    setup_protocol_table(
        table, AWS_SUMMARY_GRID_COLS, width_dxa=AWS_SUMMARY_TABLE_WIDTH_DXA, borders=True
    )

    hdr = table.rows[0]
    set_cell_text(hdr.cells[0], "Sr. No", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr.cells[1], "Tests", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr.cells[2], "")
    set_cell_text(hdr.cells[3], "Results", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(hdr.cells[4], "Limits", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for col in (3, 4):
        set_cell_no_wrap(hdr.cells[col])

    for ri, row in enumerate(rows_data):
        tr = table.rows[ri + 1]
        row_type = row.get("row_type", "standard")
        sr = row.get("sr_no", "")
        name = row.get("name", "")
        result = row.get("result", "")
        limit = row.get("limit", "")

        if row_type == "section_label":
            set_cell_text(tr.cells[0], "")
            set_cell_text(tr.cells[1], name, bold=True)
            merge_row_cells(tr, 1, 4)
        elif row_type == "group_header":
            set_cell_text(tr.cells[0], f"{sr}.", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(tr.cells[1], name)
            merge_row_cells(tr, 1, 4)
        elif row_type == "sub_test":
            set_cell_text(tr.cells[0], "")
            set_cell_text(tr.cells[1], name)
            set_cell_text(tr.cells[2], "")
            set_cell_text(tr.cells[3], result)
            set_cell_text(tr.cells[4], limit)
        else:
            set_cell_text(tr.cells[0], f"{sr}.", align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(tr.cells[1], name)
            set_cell_text(tr.cells[2], "")
            set_cell_text(tr.cells[3], result)
            set_cell_text(tr.cells[4], limit)

    finalize_protocol_table(table, AWS_SUMMARY_GRID_COLS)
    for row in table.rows:
        clear_row_height(row)

    note = context.get("compliance_note") or DEFAULT_COMPLIANCE_NOTE
    add_protocol_spacer_paragraph(
        doc,
        before=PROTOCOL_SECTION_GAP_TWIPS,
        after=PROTOCOL_SECTION_GAP_TWIPS,
        text=note,
    )


def _aws_section_content_lines(section: dict) -> list[str]:
    lines: list[str] = []

    instrument = (section.get("instrument_display") or "").strip()
    if instrument:
        lines.append(f"Instrument Used: {instrument}")

    reagent = (section.get("reagent_display") or "").strip()
    if reagent:
        lines.append(f"Reagent Used: {reagent}")

    procedure = (section.get("procedure_text") or "").strip()
    if procedure:
        lines.append("Procedure:")
        lines.extend(line for line in procedure.split("\n") if line.strip())

    readings = (section.get("readings_display") or "").strip()
    if readings:
        lines.append(f"Readings: {readings}")

    calculated = (section.get("calculated_result") or "").strip()
    if calculated:
        lines.append(f"Calculated Result: {calculated}")

    result = (section.get("result_display") or "").strip()
    if result:
        lines.append(f"Result: {result}")

    limit = (section.get("limits_display") or "").strip()
    if limit:
        lines.append(f"Acceptance criteria: {limit}")

    conclusion = (section.get("conclusion_display") or "").strip()
    if conclusion:
        lines.append(f"Conclusion: {conclusion}")

    if section.get("is_oos") and section.get("oos_acknowledged"):
        comment = (section.get("oos_ack_comment") or "").strip()
        if comment:
            lines.append(f"OOS Acknowledgement: {comment}")

    if section.get("instrument_expired_ack") or section.get("reagent_expired_ack"):
        expiry_comment = (section.get("expiry_ack_comment") or "").strip()
        if expiry_comment:
            lines.append(f"Expiry Acknowledgement: {expiry_comment}")

    analyst = section.get("analyst") or {}
    checker = section.get("checker") or {}
    analyst_name = analyst.get("name", "")
    checker_name = checker.get("name", "")
    lines.append(f"Analyzed By: {analyst_name}")
    lines.append(f"Checked By: {checker_name}")

    return lines


def _new_detail_table(doc: Document, row_count: int):
    table = doc.add_table(rows=row_count, cols=3)
    setup_protocol_table(
        table, AWS_DETAIL_GRID_COLS, borders=True, indent_dxa=AWS_HEADER_TABLE_INDENT_DXA
    )
    return table


def _write_detail_title_row(row, col0: str, title: str) -> None:
    set_cell_text(row.cells[0], col0, bold=True, line_twips=DETAIL_COMPACT_LINE_TWIPS)
    set_cell_text(row.cells[1], title, line_twips=DETAIL_COMPACT_LINE_TWIPS)
    merge_row_cells(row, 1, 2)
    for cell in row.cells:
        set_cell_margins_dxa(cell)
    clear_row_height(row)


def _write_detail_content_row(row, lines: list[str]) -> None:
    cell = row.cells[0]
    set_cell_paragraphs(cell, lines, compact=True)
    merge_row_cells(row, 0, 2)
    set_cell_valign(cell, "top")
    set_cell_margins_dxa(cell)
    clear_row_height(row)


def _build_aws_section_detail(doc: Document, section: dict, *, add_spacer_after: bool) -> None:
    section_no = section.get("section_no", "")
    name = section.get("test_name", "")

    table = _new_detail_table(doc, 2)
    _write_detail_title_row(table.rows[0], section_no, f"{name}:")
    _write_detail_content_row(table.rows[1], _aws_section_content_lines(section))
    finalize_protocol_table(table, AWS_DETAIL_GRID_COLS)
    if add_spacer_after:
        add_protocol_spacer_paragraph(doc, before=0, after=0, line=DETAIL_COMPACT_LINE_TWIPS)


def _build_aws_detail_tables(doc: Document, context: dict) -> None:
    sections = sorted(context.get("sections") or [], key=lambda s: s.get("sort_order", 0))
    for i, section in enumerate(sections):
        _build_aws_section_detail(doc, section, add_spacer_after=i < len(sections) - 1)


def build_aws_body(doc: Document, context: dict) -> None:
    clear_document_body(doc)
    _build_batch_info_table(doc, context)
    _build_aws_summary_table(doc, context)
    _build_aws_detail_tables(doc, context)
    if context.get("revision_history"):
        build_revision_history_table(
            doc,
            context.get("revision_history") or [],
            grid_cols=AWS_REVISION_GRID_COLS,
            table_width_dxa=AWS_REVISION_TABLE_WIDTH_DXA,
            indent_dxa=AWS_REVISION_TABLE_INDENT_DXA,
            label=context.get("document_no_label", "AWS NO."),
        )


def apply_aws_layout(doc: Document, context: dict) -> None:
    doc.settings.odd_and_even_pages_header_footer = False
    for section in doc.sections:
        configure_page_setup(section, DocumentType.AWS)
        wire_repeating_aws_header_footer(section, context)

    build_aws_body(doc, context)
