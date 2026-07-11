"""Document rendering engine using docxtpl and python-docx post-processing."""

from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docxtpl import DocxTemplate

from app.core.config import get_settings
from app.core.constants import DocumentType
from app.document_engine.components.header_footer import (
    add_revision_history_block,
    apply_header_footer_to_document,
)
from app.document_engine.components.aws_layout import apply_aws_layout
from app.document_engine.components.coa_layout import apply_coa_layout
from app.document_engine.components.moa_layout import apply_moa_layout
from app.document_engine.components.protocol_layout import apply_protocol_layout
from app.document_engine.components.sop_layout import apply_sop_layout
from app.document_engine.components.spec_layout import apply_spec_layout
from app.document_engine.aws_context import build_aws_context
from app.document_engine.coa_context import build_coa_context
from app.document_engine.context_builder import build_document_context
from app.document_engine.table_renderer import render_table_to_document
from app.schemas.aws_render import AwsRenderInput
from app.schemas.coa_render import CoaRenderInput
from app.schemas.product import ProductConfig


TEMPLATE_MAP = {
    DocumentType.MOA: "base_moa.docx",
    DocumentType.PROTOCOL: "base_protocol.docx",
    DocumentType.SOP: "base_sop.docx",
    DocumentType.ANNEXURE: "base_annexure.docx",
    DocumentType.SPECIFICATION: "base_specification.docx",
    DocumentType.STANDARD_FORMAT: "base_annexure.docx",
}

# Product QA documents use python-docx layouts matching reference samples.
PROGRAMMATIC_LAYOUTS = {
    DocumentType.PROTOCOL: apply_protocol_layout,
    DocumentType.MOA: apply_moa_layout,
    DocumentType.SPECIFICATION: apply_spec_layout,
    DocumentType.SOP: apply_sop_layout,
    DocumentType.COA: apply_coa_layout,
    DocumentType.AWS: apply_aws_layout,
}


class DocumentRenderer:
    def __init__(self, templates_dir: Path | None = None) -> None:
        settings = get_settings()
        self.templates_dir = templates_dir or settings.templates_dir
        self.generated_dir = settings.generated_dir

    def get_template_path(self, document_type: DocumentType) -> Path:
        filename = TEMPLATE_MAP.get(document_type)
        if not filename:
            raise ValueError(f"No template for document type: {document_type}")
        return self.templates_dir / filename

    def render(
        self,
        document_type: DocumentType,
        product_config: ProductConfig | dict,
        output_filename: str,
        **context_kwargs: Any,
    ) -> Path:
        context = build_document_context(
            product_config,
            document_type,
            **context_kwargs,
        )

        output_path = self.generated_dir / output_filename

        if document_type in PROGRAMMATIC_LAYOUTS:
            doc = DocxDocument()
            PROGRAMMATIC_LAYOUTS[document_type](doc, context)
            doc.save(str(output_path))
            return output_path

        template_path = self.get_template_path(document_type)
        if not template_path.exists():
            raise FileNotFoundError(
                f"Template not found: {template_path}. Run: python scripts/build_templates.py"
            )

        tpl = DocxTemplate(str(template_path))
        tpl.render(context)
        tpl.save(str(output_path))

        self._post_process(output_path, context, document_type)
        return output_path

    def render_coa(self, payload: CoaRenderInput, output_filename: str) -> Path:
        """Render COA from render-ready CoaRenderInput — no ProductConfig derivation."""
        context = build_coa_context(payload)
        output_path = self.generated_dir / output_filename
        doc = DocxDocument()
        apply_coa_layout(doc, context)
        doc.save(str(output_path))
        return output_path

    def render_aws(self, payload: AwsRenderInput, output_filename: str) -> Path:
        """Render AWS from render-ready AwsRenderInput — no ProductConfig derivation."""
        context = build_aws_context(payload)
        output_path = self.generated_dir / output_filename
        doc = DocxDocument()
        apply_aws_layout(doc, context)
        doc.save(str(output_path))
        return output_path

    def _post_process(
        self,
        output_path: Path,
        context: dict,
        document_type: DocumentType,
    ) -> None:
        """Apply header/footer, revision history, and dynamic tables via python-docx."""
        doc = DocxDocument(str(output_path))

        if document_type == DocumentType.SOP and context.get("revision_history"):
            add_revision_history_block(doc, context["revision_history"])

        for table_cfg in context.get("extra_tables", []):
            render_table_to_document(doc, table_cfg)

        apply_header_footer_to_document(doc, context, document_type)
        doc.save(str(output_path))
