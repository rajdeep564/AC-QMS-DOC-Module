"""Dynamic table rendering with python-docx."""

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.core.config import get_settings
from app.core.constants import DocumentType
from app.document_engine.styles import (
    apply_table_grid_borders,
    configure_page_setup,
    set_cell_text,
    table_full_width,
)


def render_table_to_document(
    doc: Document,
    table_config: dict[str, Any],
    after_paragraph=None,
) -> None:
    columns: list[str] = table_config.get("columns", [])
    rows: list[list[Any]] = table_config.get("rows", [])
    merges: list[dict[str, int]] = table_config.get("merges", [])
    table_name: str | None = table_config.get("table_name")

    settings = get_settings()

    if table_name:
        p = doc.add_paragraph()
        run = p.add_run(table_name)
        run.bold = True
        run.font.name = settings.default_font
        run.font.size = Pt(settings.default_font_size)

    if not columns:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(columns))
    table_full_width(table)
    apply_table_grid_borders(table)

    for col_idx, header in enumerate(columns):
        set_cell_text(
            table.rows[0].cells[col_idx],
            str(header),
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )

    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            if col_idx < len(columns):
                set_cell_text(table.rows[row_idx + 1].cells[col_idx], str(value) if value is not None else "")

    _apply_merges(table, merges)
    doc.add_paragraph()


def _apply_merges(table, merges: list[dict[str, int]]) -> None:
    for merge in merges:
        r1 = merge.get("row_start", 0)
        r2 = merge.get("row_end", r1)
        c1 = merge.get("col_start", 0)
        c2 = merge.get("col_end", c1)
        start_cell = table.rows[r1].cells[c1]
        end_cell = table.rows[r2].cells[c2]
        start_cell.merge(end_cell)


# Re-export for backward compatibility
def set_page_border(section, **kwargs):
    from app.document_engine.styles import set_page_border as _set

    return _set(section, **kwargs)


def configure_page_setup_legacy(section) -> None:
    configure_page_setup(section, DocumentType.SOP)
