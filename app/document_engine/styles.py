"""Document styling: margins, page borders, table borders, fonts."""

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.core.config import get_settings
from app.core.constants import DocumentType
from app.document_engine.sop_style import get_page_border_spaces, get_page_setup, load_sop_style

_cfg = load_sop_style()
_t = _cfg.tables
_p = _t.protocol
_m = _t.moa
_s = _t.specification
_sop = _t.sop
_c = _t.coa
_a = _t.aws
_sp = _cfg.spacing
_logo = _cfg.logo

# Protocol / shared table constants (from sop_style.yaml)
PROTOCOL_TABLE_WIDTH_DXA = _p.table_width_dxa
BATCH_TABLE_WIDTH_DXA = _p.batch_table_width_dxa
SUMMARY_TABLE_WIDTH_DXA = _p.summary_table_width_dxa
SIGNOFF_TABLE_WIDTH_DXA = _p.signoff_table_width_dxa
COLON_COL_WIDTH_DXA = _t.colon_col_width_dxa
BATCH_GRID_COLS = _p.batch_grid_cols
SUMMARY_GRID_COLS = _p.summary_grid_cols
SIGNOFF_GRID_COLS = _p.signoff_grid_cols
FOOTER_GRID_COLS = _t.footer_grid_cols
FOOTER_TABLE_WIDTH_DXA = _t.footer_table_width_dxa
FOOTER_TABLE_INDENT_DXA = _t.footer_table_indent_dxa
DETAIL_GRID_COLS = _p.detail_grid_cols
_TW = _t.shared_table_width_dxa

MOA_TABLE_WIDTH_DXA = _m.table_width_dxa
MOA_DETAIL_GRID_COLS = _m.detail_grid_cols
MOA_HEADER_TABLE_WIDTH_DXA = _m.header_table_width_dxa
MOA_HEADER_TABLE_INDENT_DXA = _m.header_table_indent_dxa
MOA_HEADER_GRID_COLS = _m.header_grid_cols
MOA_FOOTER_TABLE_WIDTH_DXA = _m.footer_table_width_dxa
MOA_FOOTER_GRID_COLS = _m.footer_grid_cols
MOA_REVISION_GRID_COLS = _m.revision_grid_cols
MOA_REVISION_TABLE_WIDTH_DXA = _m.revision_table_width_dxa
MOA_REVISION_TABLE_INDENT_DXA = _m.revision_table_indent_dxa

SPEC_TABLE_WIDTH_DXA = _s.table_width_dxa
SPEC_HEADER_TABLE_WIDTH_DXA = _s.header_table_width_dxa
SPEC_HEADER_TABLE_INDENT_DXA = _s.header_table_indent_dxa
SPEC_HEADER_GRID_COLS = _s.header_grid_cols
SPEC_FOOTER_TABLE_WIDTH_DXA = _s.footer_table_width_dxa
SPEC_FOOTER_GRID_COLS = _s.footer_grid_cols
SPEC_FOOTER_TABLE_INDENT_DXA = _s.footer_table_indent_dxa
SPEC_PRODUCT_GRID_COLS = _s.product_grid_cols
SPEC_PARAM_GRID_COLS = _s.param_grid_cols
SPEC_MICRO_GRID_COLS = _s.micro_grid_cols
SPEC_REVISION_GRID_COLS = _s.revision_grid_cols
SPEC_REVISION_TABLE_WIDTH_DXA = _s.revision_table_width_dxa
SPEC_REVISION_TABLE_INDENT_DXA = _s.revision_table_indent_dxa

COA_HEADER_TABLE_WIDTH_DXA = _c.header_table_width_dxa
COA_HEADER_TABLE_INDENT_DXA = _c.header_table_indent_dxa
COA_HEADER_GRID_COLS = _c.header_grid_cols
COA_FOOTER_TABLE_WIDTH_DXA = _c.footer_table_width_dxa
COA_FOOTER_GRID_COLS = _c.footer_grid_cols
COA_FOOTER_TABLE_INDENT_DXA = _c.footer_table_indent_dxa
COA_RESULTS_TABLE_WIDTH_DXA = _c.results_table_width_dxa
COA_RESULTS_GRID_COLS = _c.results_grid_cols
COA_REVISION_GRID_COLS = _c.revision_grid_cols
COA_REVISION_TABLE_WIDTH_DXA = _c.revision_table_width_dxa
COA_REVISION_TABLE_INDENT_DXA = _c.revision_table_indent_dxa

AWS_TABLE_WIDTH_DXA = _a.table_width_dxa
AWS_BATCH_TABLE_WIDTH_DXA = _a.batch_table_width_dxa
AWS_SUMMARY_TABLE_WIDTH_DXA = _a.summary_table_width_dxa
AWS_SIGNOFF_TABLE_WIDTH_DXA = _a.signoff_table_width_dxa
AWS_BATCH_GRID_COLS = _a.batch_grid_cols
AWS_SUMMARY_GRID_COLS = _a.summary_grid_cols
AWS_SIGNOFF_GRID_COLS = _a.signoff_grid_cols
AWS_DETAIL_GRID_COLS = _a.detail_grid_cols
AWS_HEADER_GRID_COLS = _a.header_grid_cols
AWS_HEADER_TABLE_WIDTH_DXA = _a.header_table_width_dxa
AWS_HEADER_TABLE_INDENT_DXA = _a.header_table_indent_dxa
AWS_FOOTER_TABLE_WIDTH_DXA = _a.footer_table_width_dxa
AWS_FOOTER_GRID_COLS = _a.footer_grid_cols
AWS_FOOTER_TABLE_INDENT_DXA = _a.footer_table_indent_dxa
AWS_REVISION_GRID_COLS = _a.aws_revision_grid_cols
AWS_REVISION_TABLE_WIDTH_DXA = _a.aws_revision_table_width_dxa
AWS_REVISION_TABLE_INDENT_DXA = _a.aws_revision_table_indent_dxa

PROTOCOL_REVISION_GRID_COLS = _p.protocol_revision_grid_cols
PROTOCOL_REVISION_TABLE_WIDTH_DXA = _p.protocol_revision_table_width_dxa
PROTOCOL_REVISION_TABLE_INDENT_DXA = _p.protocol_revision_table_indent_dxa
HEADER_GRID_COLS = _p.header_grid_cols
HEADER_TABLE_WIDTH_DXA = _p.header_table_width_dxa

PROTOCOL_SPACER_LINE_TWIPS = _sp.protocol_spacer_twips
PROTOCOL_SECTION_GAP_TWIPS = _sp.protocol_section_gap_twips
DETAIL_COMPACT_LINE_TWIPS = _sp.detail_compact_twips
DETAIL_CELL_MARGIN_TWIPS = _sp.detail_cell_margin_twips

HEADER_TABLE_INDENT_DXA = _t.header_table_indent_dxa
SIGNOFF_TABLE_INDENT_DXA = _t.signoff_table_indent_dxa
HEADER_LOGO_ROW_HEIGHT_TWIPS = _logo.row_height_twips

SOP_HEADER_GRID_COLS = _sop.header_grid_cols
SOP_HEADER_TABLE_WIDTH_DXA = _sop.header_table_width_dxa
SOP_REVISION_GRID_COLS = _sop.revision_grid_cols
SOP_REVISION_TABLE_WIDTH_DXA = _sop.revision_table_width_dxa


def style_run(run, *, bold: bool = False, size: int | None = None, font: str | None = None) -> None:
    settings = get_settings()
    page = load_sop_style().page
    run.bold = bold
    run.font.name = font or settings.default_font or page.font
    run.font.size = Pt(size or settings.default_font_size or page.font_size_pt)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
    line_twips: int | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    if line_twips is not None:
        set_paragraph_spacing_twips(p, line=line_twips)
    if text:
        run = p.add_run(text)
        style_run(run, bold=bold)


def set_cell_paragraphs(
    cell,
    lines: list[str],
    *,
    bold_first: bool = False,
    compact: bool = False,
) -> None:
    """Multiple paragraphs inside one cell. Use compact=True for tight, content-sized rows."""
    cell.text = ""
    para_idx = 0
    for i, line in enumerate(lines):
        if not line.strip():
            continue
        p = cell.paragraphs[0] if para_idx == 0 else cell.add_paragraph()
        para_idx += 1
        if compact:
            set_paragraph_spacing_twips(
                p, before=0, after=0, line=DETAIL_COMPACT_LINE_TWIPS
            )
        run = p.add_run(line)
        style_run(run, bold=(bold_first and i == 0))


def set_cell_margins_dxa(
    cell,
    *,
    top: int | None = None,
    bottom: int | None = None,
    left: int | None = None,
    right: int | None = None,
) -> None:
    sp = load_sop_style().spacing
    if top is None:
        top = sp.detail_cell_margin_twips
    if bottom is None:
        bottom = sp.detail_cell_margin_twips
    if left is None:
        left = sp.detail_cell_margin_lr_dxa
    if right is None:
        right = sp.detail_cell_margin_lr_dxa
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcMar"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def clear_row_height(row) -> None:
    """Let Word size the row to fit cell content."""
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:trHeight"))
    if existing is not None:
        tr_pr.remove(existing)


def _border_element(side: str, val: str = "single", sz: str = "4", space: str = "0", color: str = "auto"):
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    return el


def apply_cell_borders(cell, val: str | None = None, sz: str | None = None) -> None:
    borders = load_sop_style().borders
    if val is None:
        val = borders.cell_val
    if sz is None:
        sz = borders.cell_sz
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcBorders"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        tc_borders.append(_border_element(side, val, sz))
    tc_pr.append(tc_borders)


def apply_table_grid_borders(table, val: str | None = None, sz: str | None = None) -> None:
    borders = load_sop_style().borders
    if val is None:
        val = borders.cell_val
    if sz is None:
        sz = borders.cell_sz
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tbl_borders.append(_border_element(side, val, sz))
    existing = tbl_pr.find(qn("w:tblBorders"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_borders)


def set_page_border(section, *, val: str = "double", sz: str | None = None) -> None:
    style = load_sop_style()
    if sz is None:
        sz = style.borders.page_sz
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:pgBorders"))
    if existing is not None:
        sect_pr.remove(existing)

    pg_borders = OxmlElement("w:pgBorders")
    spaces = get_page_border_spaces()
    for border_name in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), val)
        border.set(qn("w:sz"), sz)
        border.set(qn("w:space"), spaces.get(border_name, "2"))
        border.set(qn("w:color"), "auto")
        pg_borders.append(border)
    sect_pr.append(pg_borders)


def _set_margins(section, margins: dict[str, float]) -> None:
    page = load_sop_style().page
    section.page_height = Inches(page.height_inches)
    section.page_width = Inches(page.width_inches)
    section.top_margin = Inches(margins["top"])
    section.bottom_margin = Inches(margins["bottom"])
    section.left_margin = Inches(margins["left"])
    section.right_margin = Inches(margins["right"])
    section.header_distance = Inches(margins["header"])
    section.footer_distance = Inches(margins["footer"])


def configure_page_setup(section, document_type: DocumentType) -> None:
    setup = get_page_setup(document_type)
    _set_margins(section, setup)
    border = setup.get("border")
    if border is not None:
        set_page_border(section, val=border)


def clear_document_body(doc) -> None:
    body = doc.element.body
    sect_pr = body.find(qn("w:sectPr"))
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)


def _get_or_create_tbl_pr(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    return tbl_pr


def table_fixed_width_dxa(table, width_dxa: int | None = PROTOCOL_TABLE_WIDTH_DXA) -> None:
    if width_dxa is None:
        return
    tbl_pr = _get_or_create_tbl_pr(table)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_pr.append(tbl_w)


def set_cell_no_wrap(cell) -> None:
    """Prevent Word from wrapping header text character-by-character in narrow cells."""
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(OxmlElement("w:noWrap"))


def set_cell_width_dxa(cell, width_dxa: int) -> None:
    """Set explicit cell width — Word ignores tblGrid without tcW on each cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:tcW"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _cell_grid_span(cell) -> int:
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return 1
    gs = tc_pr.find(qn("w:gridSpan"))
    if gs is None:
        return 1
    return int(gs.get(qn("w:val")))


def set_table_layout_fixed(table) -> None:
    tbl_pr = _get_or_create_tbl_pr(table)
    existing = tbl_pr.find(qn("w:tblLayout"))
    if existing is not None:
        tbl_pr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def apply_table_column_widths(table, col_widths: list[int]) -> None:
    """Apply tcW on every cell from gridCol widths (respects gridSpan merges).

    Uses row._tr.tc_lst instead of row.cells — python-docx exposes duplicate
  ghost cells after horizontal merge, which would misalign column indices.
    """
    from docx.table import _Cell

    for row in table.rows:
        col_idx = 0
        for tc in row._tr.tc_lst:
            cell = _Cell(tc, row)
            span = _cell_grid_span(cell)
            width = sum(col_widths[col_idx : col_idx + span])
            set_cell_width_dxa(cell, width)
            col_idx += span


def finalize_protocol_table(table, col_widths: list[int]) -> None:
    """Lock column widths after all merges — call once table content is complete."""
    set_table_layout_fixed(table)
    apply_table_column_widths(table, col_widths)


def set_table_grid_columns(table, col_widths: list[int]) -> None:
    tbl = table._tbl
    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl_pr = _get_or_create_tbl_pr(table)
    tbl.insert(tbl.index(tbl_pr) + 1, grid)


def set_grid_span(cell, span: int) -> None:
    if span <= 1:
        return
    tc_pr = cell._tc.get_or_add_tcPr()
    grid_span = OxmlElement("w:gridSpan")
    grid_span.set(qn("w:val"), str(span))
    existing = tc_pr.find(qn("w:gridSpan"))
    if existing is not None:
        tc_pr.remove(existing)
    tc_pr.append(grid_span)


def merge_row_cells(row, start: int, end: int) -> None:
    row.cells[start].merge(row.cells[end])


def set_table_style_grid(table) -> None:
    """Apply Word 'Table Grid' style — visible borders on all cells (matches reference header)."""
    try:
        table.style = "Table Grid"
    except KeyError:
        apply_table_grid_borders(table)


def set_table_jc(table, jc: str = "center") -> None:
    """Set table horizontal alignment (jc). Use 'center' for spec body tables."""
    tbl_pr = _get_or_create_tbl_pr(table)
    existing = tbl_pr.find(qn("w:jc"))
    if existing is not None:
        tbl_pr.remove(existing)
    el = OxmlElement("w:jc")
    el.set(qn("w:val"), jc)
    tbl_pr.append(el)


def set_table_indent(table, indent_dxa: int) -> None:
    tbl_pr = _get_or_create_tbl_pr(table)
    existing = tbl_pr.find(qn("w:tblInd"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)


def set_row_height(row, height_twips: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:trHeight"))
    if existing is not None:
        tr_pr.remove(existing)
    tr_height = OxmlElement("w:trHeight")
    tr_height.set(qn("w:val"), str(height_twips))
    tr_pr.append(tr_height)


def set_cell_valign(cell, align: str = "center") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:vAlign"))
    if existing is not None:
        tc_pr.remove(existing)
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), align)
    tc_pr.append(valign)


def set_paragraph_spacing_twips(
    paragraph, *, before: int = 0, after: int = 0, line: int = PROTOCOL_SPACER_LINE_TWIPS
) -> None:
    """Set paragraph spacing via OOXML twips — matches reference protocol gaps."""
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:spacing"))
    if existing is not None:
        p_pr.remove(existing)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)


def add_protocol_spacer_paragraph(
    doc,
    *,
    before: int = 0,
    after: int = 0,
    line: int = PROTOCOL_SPACER_LINE_TWIPS,
    text: str = "",
    bold: bool = False,
) -> None:
    """Spacer paragraph between protocol tables — matches GLYCINE IP.docx gaps."""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        style_run(run, bold=bold)
    set_paragraph_spacing_twips(p, before=before, after=after, line=line)


def add_header_trailing_paragraph(header) -> None:
    """Empty Header-style paragraph after header table — creates gap before batch table."""
    p = header.add_paragraph()
    try:
        p.style = "Header"
    except KeyError:
        set_paragraph_spacing_twips(p, after=0, line=240)


def add_footer_trailing_paragraph(footer) -> None:
    """Empty Footer-style paragraph after footer table — matches reference."""
    p = footer.add_paragraph()
    try:
        p.style = "Footer"
    except KeyError:
        set_paragraph_spacing_twips(p, after=0, line=240)


def add_zero_gap_paragraph(doc) -> None:
    """Tight spacer between tables — reference uses after=0, line=360."""
    add_protocol_spacer_paragraph(doc, before=0, after=0)


def setup_protocol_table(
    table,
    col_widths: list[int],
    *,
    width_dxa: int | None = PROTOCOL_TABLE_WIDTH_DXA,
    borders: bool = False,
    cell_borders: bool = False,
    grid_style: bool = False,
    indent_dxa: int | None = None,
) -> None:
    table_fixed_width_dxa(table, width_dxa)
    set_table_grid_columns(table, col_widths)
    if grid_style:
        set_table_style_grid(table)
    if borders:
        apply_table_grid_borders(table)
    if cell_borders and not grid_style:
        for row in table.rows:
            for cell in row.cells:
                apply_cell_borders(cell)
    if indent_dxa is not None:
        set_table_indent(table, indent_dxa)


# Backward compatible alias
def table_full_width(table) -> None:
    table_fixed_width_dxa(table)
